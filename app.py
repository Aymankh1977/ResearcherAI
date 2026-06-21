"""
Universal Literature Research Tool
Search PubMed, Europe PMC, OpenAlex + AI screening via Anthropic API
Configurable inclusion criteria, tiers, and Kirkpatrick levels for any research field
"""

import streamlit as st
import requests
import json
import time
import pandas as pd
from anthropic import Anthropic
from datetime import datetime
from urllib.parse import quote

st.set_page_config(
    page_title="Literature Research Tool",
    page_icon="📚",
    layout="wide"
)

# ─────────────────────────────────────────────────────────────────
# CONFIG: User-editable presets
# ─────────────────────────────────────────────────────────────────

DEFAULT_CONFIG = {
    "research_topic": "Artificial Intelligence in Dental Education and Accreditation",
    "theoretical_anchor": "the accreditation dilemma — structural tension between formal accreditation requirements and dental schools' institutional capacity to implement them (DiMaggio & Powell isomorphism; Scott's three pillars — regulative, normative, cultural-cognitive; loose coupling; Meyer & Rowan myth and ceremony)",
    "writing_style_sample": """In dental education, assessment of students is of critical value to improve their performance in clinical settings. The assessment of students' clinical performance includes various stages, where in certain cases, assessors encounter challenges in providing final grades. This study sheds a light on assessment practices in clinical settings and focuses on assessors' modulation of the whole cognitive process. The argument involves discussing critical thinking of assessors before, during, and after the event of assessment. Then, it analyzes a cognitive approach of assessment implied by assessors during students' performance. Further, it proposes a model with step-by-step approach in decision-making along with different factors, which may strongly influence final grades. Four main stages were identified for the purpose of analysis, such as pre-decision, driver, primary decision, and communication stages. Each stage was supported by literary data, along with evidences worth consideration.

Briefly, the four stages explain the flow of information during the cognitive process with conjoint factors. For example, internal and external sources are the main factors to affect a pre-decision (non-task-specific) cognitive stage in clinical settings. The driver stage starts when the assessor is present to judge the performance of specific clinical tasks of assessment. The third stage is the primary decision stage, which begins when the assessor finds or sees (interpret) students' performance according to the defined frame of reference. The resulted primary decision predicts a range of options between 'being sure' and 'uncertainty'. Then, the refinement process of decision grade, the fourth stage, is the moderation of the decision, which is affected by another set of factors, such as legal consequences, community, and patient safety to direct the decision towards grading.

Expertise is believed to be a primary internal factor, which influences clinical reasoning to update assessors' judgment capacity. Clinical reasoning consists of two types: content-dependent and context-dependent. However, in therapeutic and diagnostic reasoning, expertise generally differs among assessors. For example, in psychology, deliberate practice acts as a key to expert performance. It may further benefit clinical reasoning as that experts possess more stringent decisions than early career assessors. In the light of behavioral learning perspective theory, expertise is guided by disciplines' specific knowledge and skills, which in turn affect internal assessors' attitudes, emotions, intentions, and personalities.

Interestingly, one external factor that is found to increase the stringency of assessors' decision is the increased number of candidates during time of assessment. This could be related to the feeling of fatigue or other factors, which may be further investigated in future studies. To conclude, it is now clear that the above-mentioned internal and external factors can influence the cognitive process during the pre-decision stage to initially predict assessors' appraisals of learners' performances.""",
    "inclusion_criteria": """1. Proposes or evaluates an AI competency/curriculum framework for the field
2. Empirical study of practitioner or learner AI readiness with explicit curricular/policy recommendations (n ≥ 30)
3. Educational intervention integrating AI into curriculum with institutional scope
4. Studies accreditation alignment, quality assurance, or educational policy regarding AI
5. Foundational position paper cited by professional/accrediting bodies""",
    "exclusion_criteria": """A. Pure clinical AI diagnostic accuracy study with no educational outcome
B. LLM benchmark/MCQ accuracy test without curricular implications
C. Single-class pilot n < 30 with no institutional integration discussion
D. KAP survey without explicit curricular recommendations
E. Editorial, commentary, letter, or non-peer-reviewed material
F. Tangential AI mention (AI not a substantive component of the study)""",
    "tiers": [
        "Tier 1 – Framework",
        "Tier 2 – Needs Evidence",
        "Tier 3 – Practitioner Readiness",
        "Tier 4 – Intervention",
        "Tier 5 – Measurement",
        "Tier 6 – Review",
    ],
    "tier_descriptions": """Tier 1 – Framework: proposes/validates a competency or curriculum framework
Tier 2 – Needs Evidence: empirical evidence on learner/practitioner educational needs
Tier 3 – Practitioner Readiness: workforce preparedness, awareness, attitudes
Tier 4 – Intervention: tests an educational intervention with institutional scope
Tier 5 – Measurement: capability benchmarked against educational standards
Tier 6 – Review: scoping/systematic review or bibliometric analysis""",
    "kp_levels": ["N/A", "L1 Reaction", "L2 Learning", "L3 Behaviour", "L4 Results"],
    "kp_description": """N/A – framework, review, or non-intervention study
L1 – Reaction: attitudes, satisfaction, perceived value
L2 – Learning: knowledge/skill acquisition (objective or self-reported)
L3 – Behaviour: observed change in practice
L4 – Results: patient or institutional outcomes (Yardley & Dornan 2012)""",
    "search_presets": [
        ("Broad: AI + Dental Education", '("artificial intelligence" OR "machine learning" OR "deep learning" OR "ChatGPT" OR "large language model" OR "generative AI") AND ("dental student" OR "dental students" OR "dental education" OR "dental curriculum" OR "dental faculty")'),
        ("Narrow: + Accreditation/QA", '("artificial intelligence" OR "machine learning" OR "ChatGPT" OR "large language model") AND ("dental education" OR "dental curriculum") AND ("accreditation" OR "quality assurance" OR "competency framework")'),
        ("Educator readiness", '("artificial intelligence" OR "ChatGPT") AND ("dental faculty" OR "dental educator" OR "dental teacher") AND ("readiness" OR "attitude" OR "perception" OR "competence")'),
        ("Curriculum frameworks", '("artificial intelligence" OR "ChatGPT") AND ("dental education" OR "dental curriculum") AND ("curriculum" OR "competency" OR "framework" OR "learning outcome")'),
        ("GDC / UK standards", '("artificial intelligence" OR "ChatGPT") AND ("dental education") AND ("General Dental Council" OR "GDC" OR "United Kingdom" OR "UK")'),
    ],
}

# ─────────────────────────────────────────────────────────────────
# SESSION STATE
# ─────────────────────────────────────────────────────────────────

def init_state():
    defaults = {
        "config": DEFAULT_CONFIG.copy(),
        "results": [],
        "corpus": [],
        "search_status": "",
        "db_status": {},
        "screening_id": None,
        "synth_text": "",
        "synth_data": None,
        "themes": None,
        "synth_chat": [],
        "trends_data": None,
        "trends_papers": [],
        "appraisals": {},  # study_id -> appraisal result dict
        "elsevier_key": "",  # Optional Scopus/EMBASE API key
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

init_state()

# ─────────────────────────────────────────────────────────────────
# DATABASE SEARCH FUNCTIONS
# ─────────────────────────────────────────────────────────────────

def build_query_prompt():
    c = st.session_state.config
    return f"""You are a search-strategy expert helping construct a PubMed query for a scoping review on: {c['research_topic']}.

Inclusion criteria:
{c['inclusion_criteria']}

Exclusion criteria (these will be applied later during AI screening, NOT in the search):
{c['exclusion_criteria']}

TASK
Construct an OPTIMAL PubMed search query that maximises sensitivity (recall) for the topic. Apply these principles:

1. The search should be BROAD — it captures everything potentially relevant. Refinement happens later during screening, NOT in the query.
2. Build two-to-three CONCEPTUAL clusters joined by AND. Each cluster contains synonyms joined by OR.
   - Cluster 1: the core technology/concept (e.g., "artificial intelligence" OR "machine learning" OR "ChatGPT" OR ...)
   - Cluster 2: the population/context (e.g., "dental education" OR "dental student" OR "dental curriculum" OR ...)
   - Optional Cluster 3: only add if the topic is genuinely narrow (e.g., a specific outcome, region, or methodology that ALL eligible papers must mention in title/abstract)
3. DO NOT add restrictive clauses like "AND (accreditation OR quality)" unless the topic explicitly demands it — these clauses kill sensitivity and the concept is better detected during AI screening of abstracts.
4. Use quoted phrases for multi-word terms.
5. Use [Title/Abstract] or [Mesh] tags only if specifically needed; default to free-text search.
6. Estimate the likely yield (broad: 200-1000+, moderate: 50-200, narrow: <50).

Respond ONLY as valid JSON, no markdown:
{{
  "query": "the PubMed search string",
  "rationale": "1-2 sentence explanation of the cluster structure",
  "estimated_yield": "broad|moderate|narrow",
  "estimated_count": "rough integer estimate or range"
}}"""


def ai_generate_query():
    """Use Claude to generate an optimal PubMed query from the configured topic + criteria."""
    text, err = call_claude(build_query_prompt(), "Generate the optimal query for this scoping review.", max_tokens=1000)
    if err:
        return None, err
    try:
        cleaned = text.replace("```json", "").replace("```", "").strip()
        data = json.loads(cleaned)
        return data, None
    except Exception as e:
        return None, f"Could not parse query JSON: {e}\n\nRaw: {text[:500]}"


def pubmed_count_only(query, date_from, date_to):
    """Quickly return only the total count from PubMed without fetching items."""
    base = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils"
    try:
        s = requests.get(
            f"{base}/esearch.fcgi",
            params={
                "db": "pubmed",
                "term": query,
                "mindate": date_from,
                "maxdate": date_to,
                "retmax": 0,
                "retmode": "json",
            },
            timeout=15,
        )
        s.raise_for_status()
        sd = s.json()
        return int(sd.get("esearchresult", {}).get("count", 0)), None
    except Exception as e:
        return 0, str(e)


def epmc_count_only(query, date_from, date_to):
    """Quickly return only the total count from Europe PMC."""
    try:
        full_query = f"({query}) AND (PUB_YEAR:[{date_from} TO {date_to}])"
        r = requests.get(
            "https://www.ebi.ac.uk/europepmc/webservices/rest/search",
            params={"query": full_query, "format": "json", "pageSize": 1, "resultType": "lite"},
            timeout=15,
        )
        r.raise_for_status()
        d = r.json()
        return int(d.get("hitCount", 0)), None
    except Exception as e:
        return 0, str(e)


def openalex_count_only(query, date_from, date_to):
    """Quickly return only the total count from OpenAlex."""
    try:
        r = requests.get(
            "https://api.openalex.org/works",
            params={
                "search": query,
                "filter": f"from_publication_date:{date_from}-01-01,to_publication_date:{date_to}-12-31",
                "per-page": 1,
                "mailto": "research.tool@example.com",
            },
            timeout=15,
        )
        r.raise_for_status()
        d = r.json()
        return int(d.get("meta", {}).get("count", 0)), None
    except Exception as e:
        return 0, str(e)


# ─── Semantic Scholar ───────────────────────────────────────────
def search_semanticscholar(query, date_from, date_to, max_results):
    """Semantic Scholar API — broad coverage with AI-enhanced search, abstracts included.
    Free, no API key needed for basic use. Rate-limited ~100 req/5min without key."""
    try:
        # Cap max_results at 100 (S2 hard limit per request)
        limit = min(max_results, 100)
        r = requests.get(
            "https://api.semanticscholar.org/graph/v1/paper/search",
            params={
                "query": query,
                "limit": limit,
                "year": f"{date_from}-{date_to}",
                "fields": "title,abstract,authors,year,journal,externalIds,url",
            },
            headers={"User-Agent": "LiteratureResearchTool/1.0"},
            timeout=25,
        )
        r.raise_for_status()
        d = r.json()
        total = d.get("total", 0)
        items = []
        for p in d.get("data", []):
            authors = ", ".join([a.get("name", "") for a in (p.get("authors") or [])[:3]])
            if len((p.get("authors") or [])) > 3:
                authors += " et al."
            ext = p.get("externalIds") or {}
            pid = ext.get("DOI") or ext.get("PubMed") or p.get("paperId", "")
            items.append({
                "id": f"s2_{p.get('paperId','')}",
                "pmid": ext.get("PubMed", ""),
                "doi": ext.get("DOI", ""),
                "title": p.get("title") or "No title",
                "authors": authors,
                "journal": (p.get("journal") or {}).get("name", ""),
                "year": str(p.get("year") or ""),
                "db": "Semantic Scholar",
                "url": p.get("url") or (f"https://doi.org/{ext.get('DOI')}" if ext.get("DOI") else ""),
                "abstract": (p.get("abstract") or "")[:2000],
                "decision": "Pending",
                "tier": "",
                "kp": "N/A",
                "rationale": "",
                "confidence": "",
                "extraction": None,
            })
        return items, total
    except Exception as e:
        st.warning(f"Semantic Scholar error: {e}")
        return [], 0


def semanticscholar_count_only(query, date_from, date_to):
    try:
        r = requests.get(
            "https://api.semanticscholar.org/graph/v1/paper/search",
            params={"query": query, "limit": 1, "year": f"{date_from}-{date_to}", "fields": "title"},
            headers={"User-Agent": "LiteratureResearchTool/1.0"},
            timeout=15,
        )
        r.raise_for_status()
        return int(r.json().get("total", 0)), None
    except Exception as e:
        return 0, str(e)


# ─── Crossref ───────────────────────────────────────────────────
def search_crossref(query, date_from, date_to, max_results):
    """Crossref API — broad metadata coverage of every DOI registered with Crossref.
    Catches journals not indexed in PubMed (engineering, education-specific, regional)."""
    try:
        # Crossref caps rows at 1000 per request
        rows = min(max_results, 1000)
        r = requests.get(
            "https://api.crossref.org/works",
            params={
                "query": query,
                "rows": rows,
                "filter": f"from-pub-date:{date_from},until-pub-date:{date_to},type:journal-article",
                "select": "DOI,title,author,container-title,issued,abstract,URL",
            },
            headers={"User-Agent": "LiteratureResearchTool/1.0 (research)"},
            timeout=30,
        )
        r.raise_for_status()
        d = r.json()
        msg = d.get("message", {})
        total = msg.get("total-results", 0)
        items = []
        for w in msg.get("items", []):
            title_list = w.get("title") or []
            title = title_list[0] if title_list else "No title"
            authors_raw = w.get("author") or []
            author_names = [
                f"{(a.get('given') or '').strip()} {(a.get('family') or '').strip()}".strip()
                for a in authors_raw[:3]
            ]
            authors = ", ".join([a for a in author_names if a])
            if len(authors_raw) > 3:
                authors += " et al."
            journal = (w.get("container-title") or [""])[0]
            issued = w.get("issued", {}).get("date-parts", [[]])[0]
            year = str(issued[0]) if issued else ""
            doi = w.get("DOI", "")
            # Strip Crossref's JATS XML tags from abstract if present
            abstract = w.get("abstract") or ""
            if abstract:
                import re
                abstract = re.sub(r"<[^>]+>", " ", abstract)
                abstract = re.sub(r"\s+", " ", abstract).strip()[:2000]
            items.append({
                "id": f"cr_{doi.replace('/', '_')}" if doi else f"cr_{hash(title) & 0xfffffff:x}",
                "pmid": "",
                "doi": doi,
                "title": title,
                "authors": authors,
                "journal": journal,
                "year": year,
                "db": "Crossref",
                "url": w.get("URL") or (f"https://doi.org/{doi}" if doi else ""),
                "abstract": abstract,
                "decision": "Pending",
                "tier": "",
                "kp": "N/A",
                "rationale": "",
                "confidence": "",
                "extraction": None,
            })
        return items, total
    except Exception as e:
        st.warning(f"Crossref error: {e}")
        return [], 0


def crossref_count_only(query, date_from, date_to):
    try:
        r = requests.get(
            "https://api.crossref.org/works",
            params={
                "query": query,
                "rows": 0,
                "filter": f"from-pub-date:{date_from},until-pub-date:{date_to},type:journal-article",
            },
            headers={"User-Agent": "LiteratureResearchTool/1.0"},
            timeout=15,
        )
        r.raise_for_status()
        return int(r.json().get("message", {}).get("total-results", 0)), None
    except Exception as e:
        return 0, str(e)


# ─── ERIC (Education Resources Information Center) ──────────────
def search_eric(query, date_from, date_to, max_results):
    """ERIC API — education-specific database. Highly relevant for dental EDUCATION,
    accreditation, curriculum design, faculty development research. Free, no key."""
    try:
        rows = min(max_results, 2000)  # ERIC supports up to 2000 per request
        # Build query with date filter; ERIC uses Solr-style syntax
        full_query = f'({query}) AND publicationdateyear:[{date_from} TO {date_to}]'
        r = requests.get(
            "https://api.ies.ed.gov/eric/",
            params={
                "search": full_query,
                "format": "json",
                "rows": rows,
                "fields": "id,title,author,source,publicationdateyear,description,url",
            },
            timeout=25,
        )
        r.raise_for_status()
        d = r.json()
        resp = d.get("response", {})
        total = resp.get("numFound", 0)
        items = []
        for doc in resp.get("docs", []):
            authors_list = doc.get("author") or []
            authors = ", ".join(authors_list[:3])
            if len(authors_list) > 3:
                authors += " et al."
            eric_id = doc.get("id", "")
            items.append({
                "id": f"eric_{eric_id}",
                "pmid": "",
                "doi": "",
                "title": doc.get("title") or "No title",
                "authors": authors,
                "journal": (doc.get("source") or [""])[0] if isinstance(doc.get("source"), list) else (doc.get("source") or ""),
                "year": str(doc.get("publicationdateyear") or ""),
                "db": "ERIC",
                "url": f"https://eric.ed.gov/?id={eric_id}" if eric_id else "",
                "abstract": (doc.get("description") or "")[:2000],
                "decision": "Pending",
                "tier": "",
                "kp": "N/A",
                "rationale": "",
                "confidence": "",
                "extraction": None,
            })
        return items, total
    except Exception as e:
        st.warning(f"ERIC error: {e}")
        return [], 0


def eric_count_only(query, date_from, date_to):
    try:
        full_query = f'({query}) AND publicationdateyear:[{date_from} TO {date_to}]'
        r = requests.get(
            "https://api.ies.ed.gov/eric/",
            params={"search": full_query, "format": "json", "rows": 0},
            timeout=15,
        )
        r.raise_for_status()
        return int(r.json().get("response", {}).get("numFound", 0)), None
    except Exception as e:
        return 0, str(e)


# ─── Scopus (Elsevier — requires API key, free tier available) ──
def _get_elsevier_key():
    """Get the Elsevier API key from session state or secrets."""
    if st.session_state.get("elsevier_key"):
        return st.session_state["elsevier_key"]
    try:
        return st.secrets.get("ELSEVIER_API_KEY", "")
    except Exception:
        return ""


def search_scopus(query, date_from, date_to, max_results):
    """Scopus Search API — paid Elsevier database, free API key tier available.
    Register at https://dev.elsevier.com to obtain a key. Free tier: 20,000 req/week."""
    key = _get_elsevier_key()
    if not key:
        st.warning("Scopus skipped: no Elsevier API key. Add one in the sidebar under '🔑 API keys'.")
        return [], 0
    try:
        # Scopus caps `count` at 25 per request; paginate for larger requests
        all_items = []
        total_available = 0
        BATCH = 25
        for offset in range(0, min(max_results, 5000), BATCH):
            r = requests.get(
                "https://api.elsevier.com/content/search/scopus",
                params={
                    "query": f'TITLE-ABS-KEY({query}) AND PUBYEAR > {int(date_from)-1} AND PUBYEAR < {int(date_to)+1}',
                    "count": min(BATCH, max_results - offset),
                    "start": offset,
                },
                headers={"X-ELS-APIKey": key, "Accept": "application/json"},
                timeout=25,
            )
            if r.status_code == 401:
                st.warning("Scopus: API key rejected (check key is valid for Search API).")
                return [], 0
            if r.status_code == 429:
                st.warning("Scopus: rate limit hit. Try again later.")
                break
            r.raise_for_status()
            d = r.json()
            sr = d.get("search-results", {})
            if not total_available:
                try:
                    total_available = int(sr.get("opensearch:totalResults", 0))
                except Exception:
                    total_available = 0
            entries = sr.get("entry", [])
            if not entries or (len(entries) == 1 and entries[0].get("error")):
                break
            for e in entries:
                doi = e.get("prism:doi", "")
                title = e.get("dc:title", "No title")
                authors = e.get("dc:creator", "")
                journal = e.get("prism:publicationName", "")
                pub_date = e.get("prism:coverDate", "")
                year = pub_date.split("-")[0] if pub_date else ""
                scopus_id = e.get("dc:identifier", "").replace("SCOPUS_ID:", "")
                all_items.append({
                    "id": f"sc_{scopus_id}",
                    "pmid": e.get("pubmed-id", ""),
                    "doi": doi,
                    "title": title,
                    "authors": authors,
                    "journal": journal,
                    "year": year,
                    "db": "Scopus",
                    "url": f"https://doi.org/{doi}" if doi else "",
                    "abstract": "",  # Abstract endpoint requires separate call
                    "decision": "Pending",
                    "stage": 0,
                    "tier": "",
                    "kp": "N/A",
                    "rationale": "",
                    "confidence": "",
                    "extraction": None,
                })
            if len(entries) < BATCH:
                break
        return all_items, total_available
    except Exception as e:
        st.warning(f"Scopus error: {e}")
        return [], 0


def scopus_count_only(query, date_from, date_to):
    key = _get_elsevier_key()
    if not key:
        return 0, "No Elsevier API key"
    try:
        r = requests.get(
            "https://api.elsevier.com/content/search/scopus",
            params={
                "query": f'TITLE-ABS-KEY({query}) AND PUBYEAR > {int(date_from)-1} AND PUBYEAR < {int(date_to)+1}',
                "count": 0,
            },
            headers={"X-ELS-APIKey": key, "Accept": "application/json"},
            timeout=15,
        )
        if r.status_code == 401:
            return 0, "API key rejected"
        r.raise_for_status()
        sr = r.json().get("search-results", {})
        return int(sr.get("opensearch:totalResults", 0)), None
    except Exception as e:
        return 0, str(e)


# ─── EMBASE (Elsevier — requires PAID institutional subscription) ──
def search_embase(query, date_from, date_to, max_results):
    """EMBASE Search API — requires Elsevier API key AND institutional EMBASE subscription
    (i.e., the API key must be associated with an institution that pays for EMBASE access).
    Without subscription, the endpoint returns an empty result set even with a valid key."""
    key = _get_elsevier_key()
    if not key:
        st.warning("EMBASE skipped: no Elsevier API key. Add one in the sidebar under '🔑 API keys'.")
        return [], 0
    try:
        # EMBASE Search API endpoint
        r = requests.get(
            "https://api.elsevier.com/content/embase/article",
            params={
                "query": query,
                "date": f"{date_from}-{date_to}",
                "count": min(max_results, 100),
            },
            headers={"X-ELS-APIKey": key, "Accept": "application/json"},
            timeout=25,
        )
        if r.status_code == 401:
            st.warning("EMBASE: API key rejected or no institutional EMBASE subscription detected.")
            return [], 0
        if r.status_code == 403:
            st.warning("EMBASE: institutional subscription required. Your API key works but does not have EMBASE access.")
            return [], 0
        r.raise_for_status()
        d = r.json()
        # EMBASE response structure
        results = d.get("results", {}).get("result", []) if isinstance(d.get("results"), dict) else []
        total = int(d.get("results", {}).get("totalResults", 0)) if isinstance(d.get("results"), dict) else 0
        items = []
        for e in results:
            doi = e.get("doi", "")
            title = e.get("title", "No title")
            authors = ", ".join(e.get("authors", [])[:3])
            if len(e.get("authors", [])) > 3:
                authors += " et al."
            embase_id = e.get("embaseID", "")
            items.append({
                "id": f"em_{embase_id}",
                "pmid": e.get("pubmedID", ""),
                "doi": doi,
                "title": title,
                "authors": authors,
                "journal": e.get("source", {}).get("title", "") if isinstance(e.get("source"), dict) else "",
                "year": str(e.get("year", "")),
                "db": "EMBASE",
                "url": f"https://doi.org/{doi}" if doi else "",
                "abstract": (e.get("abstract") or "")[:2000],
                "decision": "Pending",
                "stage": 0,
                "tier": "",
                "kp": "N/A",
                "rationale": "",
                "confidence": "",
                "extraction": None,
            })
        return items, total
    except Exception as e:
        st.warning(f"EMBASE error: {e}")
        return [], 0


def embase_count_only(query, date_from, date_to):
    key = _get_elsevier_key()
    if not key:
        return 0, "No Elsevier API key"
    try:
        r = requests.get(
            "https://api.elsevier.com/content/embase/article",
            params={"query": query, "date": f"{date_from}-{date_to}", "count": 0},
            headers={"X-ELS-APIKey": key, "Accept": "application/json"},
            timeout=15,
        )
        if r.status_code in (401, 403):
            return 0, "Subscription / key required"
        r.raise_for_status()
        d = r.json()
        return int(d.get("results", {}).get("totalResults", 0) if isinstance(d.get("results"), dict) else 0), None
    except Exception as e:
        return 0, str(e)


# ─── CRITICAL APPRAISAL TOOLS ──────────────────────────────────
# Embedded checklists for the most-used appraisal frameworks.

APPRAISAL_TOOLS = {
    "JBI Cross-Sectional": {
        "description": "JBI Critical Appraisal Checklist for Analytical Cross-Sectional Studies",
        "applies_to": ["cross-sectional", "cross sectional", "survey", "questionnaire", "kap"],
        "items": [
            "Were the criteria for inclusion in the sample clearly defined?",
            "Were the study subjects and the setting described in detail?",
            "Was the exposure measured in a valid and reliable way?",
            "Were objective, standard criteria used for measurement of the condition?",
            "Were confounding factors identified?",
            "Were strategies to deal with confounding factors stated?",
            "Were the outcomes measured in a valid and reliable way?",
            "Was appropriate statistical analysis used?",
        ],
    },
    "Cochrane RoB 2.0": {
        "description": "Cochrane Risk of Bias tool 2.0 for randomised controlled trials",
        "applies_to": ["randomised controlled trial", "randomized controlled trial", "rct", "randomised trial", "randomized trial"],
        "items": [
            "Risk of bias arising from the randomisation process",
            "Risk of bias due to deviations from intended interventions",
            "Risk of bias due to missing outcome data",
            "Risk of bias in measurement of the outcome",
            "Risk of bias in selection of the reported result",
        ],
    },
    "JBI Qualitative": {
        "description": "JBI Critical Appraisal Checklist for Qualitative Research",
        "applies_to": ["qualitative", "interview", "focus group", "thematic analysis", "ethnograph", "phenomenolog", "grounded theory"],
        "items": [
            "Is there congruity between the stated philosophical perspective and the research methodology?",
            "Is there congruity between the research methodology and the research question or objectives?",
            "Is there congruity between the research methodology and the methods used to collect data?",
            "Is there congruity between the research methodology and the representation and analysis of data?",
            "Is there congruity between the research methodology and the interpretation of results?",
            "Is there a statement locating the researcher culturally or theoretically?",
            "Is the influence of the researcher on the research, and vice versa, addressed?",
            "Are participants, and their voices, adequately represented?",
            "Is the research ethical according to current criteria, or is there evidence of ethical approval by an appropriate body?",
            "Do the conclusions drawn flow from the analysis or interpretation of the data?",
        ],
    },
    "MMAT": {
        "description": "Mixed Methods Appraisal Tool (MMAT) 2018",
        "applies_to": ["mixed methods", "mixed-methods"],
        "items": [
            "Are there clear research questions?",
            "Do the collected data allow to address the research questions?",
            "Is there an adequate rationale for using a mixed methods design?",
            "Are the different components of the study effectively integrated?",
            "Are the outputs of the integration of qualitative and quantitative components adequately interpreted?",
            "Are divergences and inconsistencies between quantitative and qualitative results adequately addressed?",
            "Do the different components of the study adhere to the quality criteria of each tradition?",
        ],
    },
    "AMSTAR-2": {
        "description": "AMSTAR-2: A critical appraisal tool for systematic reviews",
        "applies_to": ["systematic review", "scoping review", "meta-analysis", "umbrella review"],
        "items": [
            "Did the research questions and inclusion criteria for the review include the components of PICO?",
            "Did the report contain an explicit statement that the review methods were established prior to conduct (protocol)?",
            "Did the review authors explain their selection of study designs for inclusion?",
            "Did the review authors use a comprehensive literature search strategy?",
            "Did the review authors perform study selection in duplicate?",
            "Did the review authors perform data extraction in duplicate?",
            "Did the review authors provide a list of excluded studies and justify the exclusions?",
            "Did the review authors describe the included studies in adequate detail?",
            "Did the review authors use a satisfactory technique for assessing risk of bias in individual studies?",
            "Did the review authors report on the sources of funding for the studies?",
            "If meta-analysis was performed, did the authors use appropriate methods for statistical combination?",
            "If meta-analysis was performed, did the authors assess the potential impact of risk of bias?",
            "Did the review authors account for risk of bias when interpreting results?",
            "Did the review authors provide a satisfactory explanation for any heterogeneity observed?",
            "If quantitative synthesis was performed, did the authors adequately investigate publication bias?",
            "Did the review authors report any potential sources of conflict of interest, including funding?",
        ],
    },
    "Opinion Papers (Modified)": {
        "description": "Modified Quality Criteria for Opinion / Position / Commentary Papers",
        "applies_to": ["opinion", "commentary", "editorial", "perspective", "position paper", "viewpoint", "narrative review", "framework", "position"],
        "items": [
            "Are the author's credentials and affiliations relevant to the topic?",
            "Was the paper published in a peer-reviewed journal of reasonable standing?",
            "Is the argument structured clearly with a defensible logical progression?",
            "Are opposing or alternative perspectives acknowledged?",
            "Is the position supported by appropriate citation of empirical evidence or theory?",
            "Is the information current at the time of writing?",
            "Is the relevance to the field clearly articulated?",
        ],
    },
}


def select_appraisal_tool(study):
    """Auto-select the most appropriate appraisal tool based on study design."""
    e = study.get("extraction") or {}
    design = (e.get("design") or "").lower()
    title = (study.get("title") or "").lower()
    abstract = (study.get("abstract") or "").lower()
    combined = f"{design} {title} {abstract[:500]}"

    for tool_name, tool_def in APPRAISAL_TOOLS.items():
        for kw in tool_def["applies_to"]:
            if kw in combined:
                return tool_name
    return "JBI Cross-Sectional"  # default


def build_appraisal_prompt(tool_name):
    tool = APPRAISAL_TOOLS[tool_name]
    items_block = "\n".join([f"  {i+1}. {item}" for i, item in enumerate(tool["items"])])
    return f"""You are conducting critical appraisal of a research study using the {tool_name} checklist.

CHECKLIST: {tool['description']}

ITEMS TO SCORE:
{items_block}

TASK
For each checklist item, return one of: "Yes", "No", "Unclear", "NA". Use the available information; if the study does not provide enough detail to judge, return "Unclear" — not "No".

Then provide:
- An OVERALL quality rating: "High" (≥80% Yes among scorable items), "Moderate" (50–79% Yes), or "Low" (<50% Yes OR critical methodological flaw)
- A short (1–2 sentence) RATIONALE for the rating
- A recommendation: "Include" if quality is High or Moderate AND the study is methodologically sound; "Exclude" if quality is Low or there are critical flaws

Respond ONLY as valid JSON, no markdown:
{{
  "tool": "{tool_name}",
  "item_scores": [
    {{ "item": "item text shortened", "rating": "Yes|No|Unclear|NA", "note": "1-line justification" }}
  ],
  "overall_rating": "High|Moderate|Low",
  "rationale": "1-2 sentence explanation",
  "recommendation": "Include|Exclude",
  "exclusion_reason": "(only if recommendation=Exclude) brief reason"
}}"""


def appraise_study(study, tool_name=None):
    """Run AI critical appraisal of a study using the appropriate tool."""
    if tool_name is None:
        tool_name = select_appraisal_tool(study)
    e = study.get("extraction") or {}
    user_msg = f"""STUDY TO APPRAISE:

Title: {study.get('title','')}
Authors: {study.get('authors','')}
Year: {study.get('year','')}
Journal: {study.get('journal','')}

Abstract: {(study.get('abstract') or '')[:2000]}

Extracted information:
- Country/Setting: {e.get('country','')}
- Design: {e.get('design','')}
- Sample: {e.get('sample','')}
- Educational Outcome: {e.get('educationalOutcome','')}
- Key Finding: {e.get('keyFinding','')}
- Main Limitation: {e.get('mainLimitation','')}
"""
    text, err = call_claude(build_appraisal_prompt(tool_name), user_msg, max_tokens=2000)
    if err:
        return None, err
    try:
        cleaned = text.replace("```json", "").replace("```", "").strip()
        return json.loads(cleaned), None
    except Exception as ex:
        return None, f"Parse error: {ex}\nRaw: {text[:500]}"


# ─── JOURNAL QUALITY CHECK ─────────────────────────────────────

PREDATORY_PATTERNS = [
    "omics international", "scientific research publishing",
    "academic journals", "international journal of recent",
    "research journal of agriculture", "world academy of",
    "scholars journal of", "open access journals press",
    "global research", "european journal of academic",
]

QUALITY_PUBLISHERS = [
    "elsevier", "springer", "wiley", "bmj", "lancet", "nature", "science",
    "sage", "taylor & francis", "routledge", "oxford university press",
    "cambridge university press", "jama", "nejm", "plos",
    "cochrane", "biomed central", "frontiers in",
    "journal of dental", "european journal of dental",
    "international dental journal", "british dental",
    "journal of dentistry",
    "medical education", "academic medicine", "advances in health",
    "bmc ", "jmir ", "journal of medical internet",
]


def check_journal_quality(journal_name):
    """Return a quality assessment based on journal name patterns."""
    if not journal_name:
        return {"flag": "warning", "reason": "Journal name missing", "advice": "Verify before inclusion"}
    j = journal_name.lower()

    for pat in PREDATORY_PATTERNS:
        if pat in j:
            return {
                "flag": "critical",
                "reason": f"Matches known predatory pattern: '{pat}'",
                "advice": "Recommend EXCLUSION — verify publisher independently",
            }

    for pub in QUALITY_PUBLISHERS:
        if pub in j:
            return {
                "flag": "good",
                "reason": f"Reputable publisher/journal ({pub})",
                "advice": "Journal quality acceptable",
            }

    if ("international journal of" in j and len(j) < 50) or ("world journal" in j) or ("global journal" in j):
        return {
            "flag": "warning",
            "reason": "Generic title pattern — verify indexing (MEDLINE / Scopus / DOAJ)",
            "advice": "Check impact factor and indexing before inclusion",
        }

    return {
        "flag": "warning",
        "reason": "Journal not on quality whitelist — verify manually",
        "advice": "Check journal indexing and peer-review process",
    }


# ─── MANCHESTER-HARVARD REFERENCING ────────────────────────────

def _format_authors_harvard(authors_str, max_authors=20):
    """Format authors for Manchester-Harvard reference list.
    'Smith J, Jones K, Brown L' -> 'Smith, J., Jones, K. and Brown, L.'"""
    if not authors_str:
        return "Anonymous"
    has_etal = "et al" in authors_str.lower()
    raw = authors_str.replace(" et al.", "").replace(" et al", "")
    parts = [a.strip() for a in raw.split(",") if a.strip()]
    formatted = []
    for p in parts[:max_authors]:
        tokens = p.split()
        if not tokens:
            continue
        # PubMed format: "Smith JA" (surname + initials run together)
        if len(tokens) >= 2 and len(tokens[-1]) <= 4 and tokens[-1].isupper():
            surname = " ".join(tokens[:-1])
            initials = ".".join(list(tokens[-1])) + "."
            formatted.append(f"{surname}, {initials}")
        elif len(tokens) >= 2:
            # "FirstName Surname" — flip to "Surname, F."
            surname = tokens[-1]
            initials = ".".join([t[0] for t in tokens[:-1] if t]) + "."
            formatted.append(f"{surname}, {initials}")
        else:
            formatted.append(p)
    if has_etal or len(parts) > max_authors:
        if len(formatted) > 1:
            return ", ".join(formatted[:-1]) + " et al."
        return (formatted[0] if formatted else "Anonymous") + " et al."
    if not formatted:
        return "Anonymous"
    if len(formatted) == 1:
        return formatted[0]
    if len(formatted) == 2:
        return f"{formatted[0]} and {formatted[1]}"
    return ", ".join(formatted[:-1]) + " and " + formatted[-1]


def _extract_surname(authors_str):
    """Extract the first author's surname for in-text citations."""
    if not authors_str:
        return "Anonymous"
    first = authors_str.split(",")[0].strip()
    tokens = first.split()
    if len(tokens) >= 2 and len(tokens[-1]) <= 4 and tokens[-1].isupper():
        return " ".join(tokens[:-1])
    if len(tokens) >= 2:
        return tokens[-1]
    return first


def format_intext_citation(study):
    """Manchester-Harvard in-text citation: (Smith, 2023) or (Smith et al., 2023)"""
    authors_str = study.get("authors", "")
    year = study.get("year", "n.d.")
    surname = _extract_surname(authors_str)
    parts = [a.strip() for a in authors_str.split(",") if a.strip()]
    has_etal = "et al" in authors_str.lower() or len(parts) > 3
    if has_etal:
        return f"({surname} et al., {year})"
    if len(parts) == 2:
        second_full = parts[1].strip()
        tokens = second_full.split()
        if len(tokens) >= 2 and len(tokens[-1]) <= 4 and tokens[-1].isupper():
            surname2 = " ".join(tokens[:-1])
        elif len(tokens) >= 2:
            surname2 = tokens[-1]
        else:
            surname2 = second_full
        return f"({surname} and {surname2}, {year})"
    if len(parts) == 3:
        return f"({surname} et al., {year})"
    return f"({surname}, {year})"


def format_reference_harvard(study):
    """Format a single study as a Manchester-Harvard reference list entry."""
    authors = _format_authors_harvard(study.get("authors", ""))
    year = study.get("year", "n.d.")
    title = (study.get("title", "") or "").rstrip(".")
    journal = study.get("journal", "")
    doi = (study.get("doi") or "").strip()

    ref = f"{authors} ({year}) '{title}'"
    if journal:
        ref += f", *{journal}*"
    ref += "."
    if doi:
        doi_clean = doi.replace("https://doi.org/", "").replace("http://doi.org/", "").replace("doi:", "").strip()
        ref += f" Available at: https://doi.org/{doi_clean}."
    elif study.get("pmid"):
        ref += f" Available at: https://pubmed.ncbi.nlm.nih.gov/{study['pmid']}/."
    elif study.get("url"):
        ref += f" Available at: {study['url']}."
    return ref


def build_reference_list(corpus):
    """Build a sorted Manchester-Harvard reference list from the corpus."""
    refs = []
    for c in corpus:
        ref = format_reference_harvard(c)
        surname = _extract_surname(c.get("authors", ""))
        refs.append((surname.lower(), c.get("year", ""), ref))
    refs.sort(key=lambda x: (x[0], x[1]))
    return [r[2] for r in refs]


def search_pubmed(query, date_from, date_to, max_results):
    """PubMed search with pagination support — fetches up to max_results IDs,
    then batches esummary/efetch calls in chunks of 200 to respect URL limits."""
    base = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils"
    try:
        # Step 1: get ALL matching PMIDs up to max_results
        s = requests.get(
            f"{base}/esearch.fcgi",
            params={
                "db": "pubmed",
                "term": query,
                "mindate": date_from,
                "maxdate": date_to,
                "retmax": max_results,
                "retmode": "json",
            },
            timeout=20,
        )
        s.raise_for_status()
        sd = s.json()
        ids = sd.get("esearchresult", {}).get("idlist", [])
        total_available = int(sd.get("esearchresult", {}).get("count", 0))
        if not ids:
            return [], total_available

        # Step 2: batch-fetch metadata in chunks of 200 (URL length safety)
        BATCH = 200
        summaries = {}
        abstracts = {}
        import re

        for batch_start in range(0, len(ids), BATCH):
            batch_ids = ids[batch_start : batch_start + BATCH]
            id_str = ",".join(batch_ids)

            # Summaries
            try:
                f = requests.post(
                    f"{base}/esummary.fcgi",
                    data={"db": "pubmed", "id": id_str, "retmode": "json"},
                    timeout=25,
                )
                f.raise_for_status()
                fd = f.json()
                summaries.update(fd.get("result", {}))
            except Exception as e:
                st.warning(f"PubMed summary batch failed at offset {batch_start}: {e}")
                continue

            # Abstracts (best-effort)
            try:
                ab = requests.post(
                    f"{base}/efetch.fcgi",
                    data={"db": "pubmed", "id": id_str, "rettype": "abstract", "retmode": "xml"},
                    timeout=30,
                )
                ab.raise_for_status()
                xml = ab.text
                articles = re.split(r"<PubmedArticle>", xml)
                for art in articles[1:]:
                    pmid_m = re.search(r"<PMID[^>]*>(\d+)</PMID>", art)
                    abs_m = re.findall(r"<AbstractText[^>]*>(.*?)</AbstractText>", art, re.DOTALL)
                    if pmid_m:
                        abstracts[pmid_m.group(1)] = " ".join(
                            [re.sub(r"<[^>]+>", "", a).strip() for a in abs_m]
                        )[:2000]
            except Exception:
                pass  # Abstracts are nice-to-have

        items = []
        for pmid in ids:
            d = summaries.get(pmid)
            if not d or pmid == "uids":
                continue
            title = (d.get("title") or "No title").replace("<", " ").replace(">", " ")
            authors = ", ".join([a.get("name", "") for a in d.get("authors", [])[:3]])
            if len(d.get("authors", [])) > 3:
                authors += " et al."
            items.append({
                "id": f"pm_{pmid}",
                "pmid": pmid,
                "title": title,
                "authors": authors,
                "journal": d.get("fulljournalname") or d.get("source", ""),
                "year": (d.get("pubdate") or "").split(" ")[0],
                "db": "PubMed",
                "url": f"https://pubmed.ncbi.nlm.nih.gov/{pmid}/",
                "abstract": abstracts.get(pmid, ""),
                "decision": "Pending",
                "tier": "",
                "kp": "N/A",
                "rationale": "",
                "confidence": "",
                "extraction": None,
            })
        return items, total_available
    except Exception as e:
        st.warning(f"PubMed error: {e}")
        return [], 0


def search_epmc(query, date_from, date_to, max_results):
    try:
        # Europe PMC supports PubMed-style queries but works best with PUB_YEAR filter
        epmc_query = f"({query}) AND (PUB_YEAR:[{date_from} TO {date_to}])"
        r = requests.get(
            "https://www.ebi.ac.uk/europepmc/webservices/rest/search",
            params={
                "query": epmc_query,
                "pageSize": max_results,
                "format": "json",
                "resultType": "core",
            },
            timeout=20,
        )
        r.raise_for_status()
        d = r.json()
        items = []
        for a in d.get("resultList", {}).get("result", []):
            authors_list = a.get("authorList", {}).get("author", [])
            authors = ", ".join([(u.get("fullName") or u.get("lastName") or "") for u in authors_list[:3]])
            if len(authors_list) > 3:
                authors += " et al."
            pmid = a.get("pmid") or a.get("id", "")
            items.append({
                "id": f"ep_{pmid}",
                "pmid": pmid,
                "title": (a.get("title") or "No title").replace("<", " ").replace(">", " "),
                "authors": authors,
                "journal": a.get("journalTitle", ""),
                "year": str(a.get("pubYear", "")),
                "db": "Europe PMC",
                "url": f"https://pubmed.ncbi.nlm.nih.gov/{a.get('pmid')}/" if a.get("pmid") else f"https://europepmc.org/article/{a.get('source','MED')}/{a.get('id')}",
                "abstract": a.get("abstractText", ""),
                "decision": "Pending",
                "tier": "",
                "kp": "N/A",
                "rationale": "",
                "confidence": "",
                "extraction": None,
            })
        return items, d.get("hitCount", len(items))
    except Exception as e:
        st.warning(f"Europe PMC error: {e}")
        return [], 0


def search_openalex(query, date_from, date_to, max_results):
    try:
        # OpenAlex works better with the 'search' parameter (full-text) than 'default.search' filter
        # Strip Boolean operators that OpenAlex doesn't handle well
        clean_query = query.replace(' AND ', ' ').replace(' OR ', ' ').replace('(', '').replace(')', '').replace('"', '')
        r = requests.get(
            "https://api.openalex.org/works",
            params={
                "search": clean_query,
                "filter": f"from_publication_date:{date_from}-01-01,to_publication_date:{date_to}-12-31",
                "per-page": max_results,
                "select": "id,title,authorships,publication_year,primary_location,doi,abstract_inverted_index",
                "mailto": "research.tool@example.com",
            },
            timeout=20,
        )
        r.raise_for_status()
        d = r.json()
        items = []
        for w in d.get("results", []):
            wid = w["id"].split("/")[-1]
            authors_list = w.get("authorships", [])
            authors = ", ".join([(a.get("author", {}).get("display_name") or "") for a in authors_list[:3] if a.get("author")])
            if len(authors_list) > 3:
                authors += " et al."
            doi = w.get("doi", "")

            # Reconstruct abstract from OpenAlex inverted index
            abstract = ""
            inv = w.get("abstract_inverted_index")
            if inv:
                word_positions = []
                for word, positions in inv.items():
                    for pos in positions:
                        word_positions.append((pos, word))
                word_positions.sort()
                abstract = " ".join([w for _, w in word_positions])[:2000]

            items.append({
                "id": f"oa_{wid}",
                "pmid": doi.replace("https://doi.org/", "") if doi else wid,
                "title": w.get("title") or "No title",
                "authors": authors,
                "journal": (w.get("primary_location") or {}).get("source", {}).get("display_name", "") if (w.get("primary_location") or {}).get("source") else "",
                "year": str(w.get("publication_year", "")),
                "db": "OpenAlex",
                "url": doi if doi else f"https://openalex.org/{wid}",
                "abstract": abstract,
                "decision": "Pending",
                "tier": "",
                "kp": "N/A",
                "rationale": "",
                "confidence": "",
                "extraction": None,
            })
        return items, d.get("meta", {}).get("count", len(items))
    except Exception as e:
        st.warning(f"OpenAlex error: {e}")
        return [], 0


def dedup_results(items):
    seen = set()
    out = []
    for it in items:
        k = "".join(c for c in (it["title"] or "").lower() if c.isalnum())[:50]
        if not k or k in seen:
            continue
        seen.add(k)
        out.append(it)
    return out


# ─────────────────────────────────────────────────────────────────
# CLAUDE AI FUNCTIONS
# ─────────────────────────────────────────────────────────────────

def get_anthropic_client():
    """Read API key from secrets or environment."""
    key = None
    try:
        key = st.secrets.get("ANTHROPIC_API_KEY")
    except Exception:
        pass
    if not key:
        import os
        key = os.environ.get("ANTHROPIC_API_KEY")
    if not key:
        return None
    return Anthropic(api_key=key)


def build_screen_system_prompt():
    c = st.session_state.config
    tiers_list = " | ".join(c["tiers"])
    kp_list = " | ".join(c["kp_levels"])
    return f"""You are a systematic review screener for a scoping review on: {c['research_topic']}.

Theoretical anchor: {c['theoretical_anchor']}

INCLUSION CRITERIA (must meet at least one):
{c['inclusion_criteria']}

EXCLUSION CRITERIA (any one excludes):
{c['exclusion_criteria']}

TIER DEFINITIONS:
{c['tier_descriptions']}

KIRKPATRICK LEVELS:
{c['kp_description']}

Respond ONLY as valid JSON with no markdown, no preamble:
{{"decision":"Include|Exclude|Maybe","tier":"<one of: {tiers_list}>","kp":"<one of: {kp_list}>","rationale":"One sentence citing which criterion drove the decision.","confidence":"high|medium|low"}}"""


def build_extract_system_prompt():
    c = st.session_state.config
    return f"""You are a data extractor for a scoping review on: {c['research_topic']}.
Based on the article title, journal, tier and year, infer the most plausible study characteristics.
For any field where inference is uncertain, prefix with "?".
Respond ONLY as valid JSON, no markdown:
{{"country":"","setting":"","design":"","sample":"","aiTool":"","educationalOutcome":"","keyFinding":"one sentence","mainLimitation":"one sentence","relevance":"direct|partial|indirect","notes":""}}"""


def build_synth_system_prompt(use_themes=False, themes=None):
    c = st.session_state.config
    style_section = ""
    if c.get("writing_style_sample", "").strip():
        style_section = f"""

═══════════════════════════════════════════════
WRITING STYLE — MIMIC THIS VOICE EXACTLY
═══════════════════════════════════════════════
The following passages illustrate the author's writing voice. Match this style precisely in every narrative section you produce: sentence rhythm, vocabulary, hedging patterns, paragraph openers, use of connective adverbs (Moreover, Therefore, However, Interestingly, Briefly, In fact, Consequently), academic third-person tone, and the characteristic move of defining a concept then giving an example.

STYLE SAMPLE:
\"\"\"
{c['writing_style_sample']}
\"\"\"

Distinctive features to replicate:
- Opens sentences with connective adverbs: "Moreover,", "Therefore,", "However,", "Interestingly,", "In fact,", "Briefly,", "Consequently,", "Nevertheless,"
- Hedging language throughout: "may," "tend to," "could be related to," "appears to," "It has been reported that," "the findings suggest"
- Scholarly third-person voice; no first person; no "we"
- Defines a concept, then illustrates with "For example,"
- Closes paragraphs with measured synthesis using "Therefore," or "It would seem that" or "Such findings suggest" — NOT with absolute claims
- Medium-long sentences with multiple clauses linked by commas
- Avoids markdown bold within narrative prose
═══════════════════════════════════════════════

╔═══════════════════════════════════════════════╗
║ HARD RULE: NO ASSERTIVE OR ABSOLUTE PHRASING  ║
╚═══════════════════════════════════════════════╝
You MUST avoid all forms of definitive, absolute, or assertive claims. The narrative must remain measured, tentative, and hedged at every turn.

PROHIBITED PHRASES (do NOT use any of these):
- "This proves..." / "It proves..." / "Proves that..."
- "Clearly demonstrates..." / "Clearly shows..."
- "Definitively..." / "Definitely..." / "Undoubtedly..." / "Without doubt..."
- "It is certain that..." / "Certainly..."
- "Must..." (in claims about what the field/policy must do)
- "Always..." / "Never..." / "All studies..."
- "The evidence shows..." (use "The evidence suggests..." instead)
- "This establishes..." / "This confirms..."
- "Strongly indicates..." (use "appears to indicate..." instead)
- "It is now clear that..." (use "Such findings suggest..." instead)

REQUIRED HEDGING (use these patterns instead):
- "The findings suggest that..." / "The evidence appears to indicate..."
- "It would seem that..." / "There is some indication that..."
- "The available data point toward..." / "These observations may reflect..."
- "Such patterns could be interpreted as..." / "One possible reading of..."
- "Within the limits of the included studies..." / "On the basis of this corpus..."
- "Tentatively..." / "Provisionally..." / "Cautiously..."

Even in the Conclusion, claims must be framed tentatively. Replace "To conclude, it is now clear that..." with "To conclude, the synthesis tentatively suggests that..." or "On balance, the corpus appears to support..."
"""

    # Choose organizing structure: emergent themes OR pre-defined tiers
    if use_themes and themes:
        theme_listing = "\n".join([
            f"  • Theme {i+1}: {t.get('name','')} — {t.get('definition','')[:150]}"
            for i, t in enumerate(themes)
        ])
        organizing_block = f"""
ORGANIZING STRUCTURE — EMERGENT THEMES
The corpus has been organized into the following data-driven themes (NOT pre-defined tiers):

{theme_listing}

When discussing tier-level patterns or distributions, instead refer to THEMES by name (or "Theme 1", "Theme 2", etc.). Theme-based synthesis replaces tier-based synthesis throughout.

Produce narrative for NINE sections, separated by the literal marker `===SECTION===` on its own line.

RESULTS SECTIONS (sections 1–5, factual interpretation, 100–180 words each)

1. CORPUS OVERVIEW — Interpret what the corpus size, year range, database spread, and theme mix tentatively suggest about the maturity and shape of the evidence base.

2. THEMATIC LANDSCAPE — Interpret which themes appear to carry the most weight (most supporting studies), which seem less developed, and what this may imply for the strength of the evidence. Refer to themes by name explicitly.

3. OUTCOME EVIDENCE PROFILE — Discuss what kinds of outcomes the corpus appears to measure (attitudes, knowledge, skill, behaviour, results) and what this may mean for the strength of claims that can be defended.

4. GEOGRAPHIC AND CONTEXTUAL VARIATION — Discuss how the spread (or absence) of geographic variation could affect generalisability and the application of the theoretical anchor.

5. CROSS-THEME PATTERNS — Identify 2-3 patterns that appear across themes. Make explicit cross-references between theme names (e.g., "When Theme 1 is read in conjunction with Theme 3...") so the Discussion can build on them.
"""
    else:
        organizing_block = """
ORGANIZING STRUCTURE — TIERS
Produce narrative for NINE sections, separated by the literal marker `===SECTION===` on its own line.

RESULTS SECTIONS (sections 1–5, factual interpretation, 100–180 words each)

1. CORPUS OVERVIEW — Interpret what the corpus size, year range, database spread, and tier mix tentatively suggest about the maturity and shape of the evidence base. Hedge throughout.

2. TIER DISTRIBUTION AND EVIDENCE WEIGHT — Interpret which tiers appear to carry the most weight, which seem under-represented, and what this may imply for the strength of the evidence. Refer to specific tier names explicitly.

3. KIRKPATRICK OUTCOME CEILING — Interpret the highest outcome level reached across the intervention studies and what this may mean for policy claims that could or could not be tentatively supported.

4. GEOGRAPHIC AND CONTEXTUAL VARIATION — Discuss how the spread (or absence) of geographic variation could affect generalisability and the application of the theoretical anchor.

5. CROSS-TIER PATTERNS — Identify 2-3 patterns that appear across tiers. Make explicit cross-references between tier numbers so the Discussion can build on them.
"""

    return f"""You are a PhD-level synthesis analyst writing in a specific author's voice for a scoping review on: {c['research_topic']}.

Theoretical anchor: {c['theoretical_anchor']}
{style_section}
TASK
You will receive (1) pre-computed statistical tables about a corpus, and (2) a list of included studies. Your job is to produce ONLY the NARRATIVE TEXT for each numbered section described below. Do NOT regenerate the tables — they are inserted programmatically. Do NOT add markdown headers (the section headings are inserted by the system).
{organizing_block}

DISCUSSION SECTIONS (sections 6–8, building toward an argument, 200–300 words each)

6. COMPARISON WITH AVAILABLE LITERATURE — Position the findings against what comparable scoping reviews, position papers, and frameworks in the wider field have reported. Note convergences and divergences. Where the corpus appears to extend or contradict prior reviews, say so tentatively. Reference the theoretical anchor explicitly. If specific comparator works are not visible in the corpus, frame the comparison at the level of general patterns reported elsewhere in the discipline.

7. INTEGRATION ACROSS {"THEMES" if use_themes and themes else "TIERS"} — Build a sustained argument that explicitly links the findings from the Results section. Use phrases like "When read in conjunction with...", "Taken together, {"Themes" if use_themes else "Tiers"} 2 and 3 appear to suggest...", "Such a pattern may be interpreted in light of the theoretical anchor...". This section must reference specific {"theme names" if use_themes and themes else "tier numbers"} and pull the threads together.

8. ARGUMENT AND IMPLICATIONS — Develop the strongest defensible argument the corpus appears to support, while remaining hedged throughout. Distinguish what the corpus may tentatively support from what it does not. Anticipate counter-positions and address them measuredly. Close this section with a clear (but hedged) claim about what the synthesis appears to contribute.

CONCLUSION (section 9, 120–180 words)

9. CONCLUSION — A closing paragraph framed measuredly. Open with phrases such as "To conclude, the synthesis tentatively suggests that...", "On balance, the corpus appears to indicate...", or "Within the limits of the included studies, the present review provisionally points toward...". Recap the central argument from Section 8 in hedged form, restate the most important research gap, and gesture toward future work.

CRITICAL RULES
- Output narrative ONLY. No markdown headers, no tables, no bullet lists.
- Use the literal marker `===SECTION===` between sections — exactly nine sections.
- Mimic the writing style sample precisely.
- Reference specific {"theme names" if use_themes and themes else "tier names and Kirkpatrick levels"} explicitly in Sections 5, 7, and 8.
- Avoid every form of assertive phrasing listed above. Hedge throughout.
- Do not invent studies or findings beyond what is in the corpus list.
- Do not fabricate citations to specific works that are not in the corpus.

REFERENCING STYLE — MANCHESTER-HARVARD
- All in-text citations must use Manchester-Harvard format: (Smith, 2023), (Smith and Jones, 2023), or (Smith et al., 2023) for three or more authors.
- Multiple citations in same parenthesis: (Smith, 2023; Jones, 2024).
- Author-prominent citations: Smith (2023) reported... or Smith and Jones (2023) found... or Smith et al. (2023) argued...
- Use citations whenever referencing a study from the corpus. Each substantive claim should be supported by 1-3 citations.
- Do NOT use numbered citations [1], [2] — use author-year only.
- The corpus listing below contains the studies you may cite. Only cite from this list."""


def compute_corpus_stats(corpus, config):
    """Compute structured statistics from the corpus for tables."""
    if not corpus:
        return {}

    n = len(corpus)
    years = [int(c["year"]) for c in corpus if c.get("year", "").isdigit()]
    year_range = f"{min(years)}–{max(years)}" if years else "—"

    # Tier distribution
    tier_counts = {}
    for t in config["tiers"]:
        tier_counts[t] = sum(1 for c in corpus if c.get("tier") == t)
    untiered = sum(1 for c in corpus if not c.get("tier"))
    if untiered:
        tier_counts["(Untiered)"] = untiered

    # Kirkpatrick distribution
    kp_counts = {}
    for k in config["kp_levels"]:
        kp_counts[k] = sum(1 for c in corpus if c.get("kp") == k)

    # Database distribution
    db_counts = {}
    for c in corpus:
        db_counts[c["db"]] = db_counts.get(c["db"], 0) + 1

    # Geographic distribution from extractions
    geo_counts = {}
    design_counts = {}
    relevance_counts = {"direct": 0, "partial": 0, "indirect": 0}
    for c in corpus:
        e = c.get("extraction") or {}
        country = (e.get("country") or "").strip().lstrip("?").strip()
        if country:
            geo_counts[country] = geo_counts.get(country, 0) + 1
        design = (e.get("design") or "").strip().lstrip("?").strip()
        if design:
            design_counts[design] = design_counts.get(design, 0) + 1
        rel = (e.get("relevance") or "").strip().lower()
        if rel in relevance_counts:
            relevance_counts[rel] += 1

    return {
        "n": n,
        "year_range": year_range,
        "tier_counts": tier_counts,
        "kp_counts": kp_counts,
        "db_counts": db_counts,
        "geo_counts": geo_counts,
        "design_counts": design_counts,
        "relevance_counts": relevance_counts,
    }


def build_tier_table(stats, config):
    """Render a tier distribution table as markdown + a list of top studies per tier."""
    rows = ["| Tier | n | % of corpus |", "|---|---|---|"]
    total = stats["n"]
    for tier, count in stats["tier_counts"].items():
        if count == 0:
            continue
        pct = round(100 * count / total, 1) if total else 0
        rows.append(f"| {tier} | {count} | {pct}% |")
    return "\n".join(rows)


def build_kp_table(stats):
    rows = ["| Kirkpatrick level | n | % of corpus |", "|---|---|---|"]
    total = stats["n"]
    for kp, count in stats["kp_counts"].items():
        if count == 0:
            continue
        pct = round(100 * count / total, 1) if total else 0
        rows.append(f"| {kp} | {count} | {pct}% |")
    return "\n".join(rows)


def build_geo_table(stats):
    if not stats["geo_counts"]:
        return None
    rows = ["| Country / setting | n |", "|---|---|"]
    sorted_geo = sorted(stats["geo_counts"].items(), key=lambda x: -x[1])
    for country, count in sorted_geo[:15]:
        rows.append(f"| {country} | {count} |")
    return "\n".join(rows)


def build_db_table(stats):
    rows = ["| Database | n studies sourced |", "|---|---|"]
    for db, count in sorted(stats["db_counts"].items(), key=lambda x: -x[1]):
        rows.append(f"| {db} | {count} |")
    return "\n".join(rows)


def _short_authors(authors):
    """Convert 'Smith J, Jones K et al.' to 'Smith et al.' or 'Smith and Jones' if 2 authors."""
    if not authors:
        return "—"
    a = authors.strip()
    # If already abbreviated with et al
    if "et al" in a.lower():
        first = a.split(",")[0].strip().split(" ")[0]
        return f"{first} et al."
    parts = [p.strip() for p in a.split(",") if p.strip()]
    if len(parts) == 1:
        return parts[0].split(" ")[0]
    elif len(parts) == 2:
        return f"{parts[0].split(' ')[0]} & {parts[1].split(' ')[0]}"
    else:
        return f"{parts[0].split(' ')[0]} et al."


def _clean_field(v, max_len=80):
    """Clean a field value for table display — handle ?, None, empty, and truncate."""
    if v is None:
        return "—"
    s = str(v).strip().lstrip("?").strip()
    if not s:
        return "—"
    if len(s) > max_len:
        return s[:max_len - 1].rstrip() + "…"
    return s


def build_overview_table_df(corpus, config):
    """Build a pandas DataFrame of included studies for Streamlit display.
    Sorted by tier order then year ascending."""
    if not corpus:
        return None

    tier_order = {t: i for i, t in enumerate(config["tiers"])}
    tier_order[""] = 999  # untiered at end

    def sort_key(c):
        return (tier_order.get(c.get("tier", ""), 999), c.get("year", ""))

    sorted_corpus = sorted(corpus, key=sort_key)

    rows = []
    for i, c in enumerate(sorted_corpus, 1):
        e = c.get("extraction") or {}
        rows.append({
            "#": i,
            "Tier": c.get("tier", "—") or "—",
            "Study": f"{_short_authors(c.get('authors',''))} ({c.get('year','')})",
            "Country/Setting": _clean_field(e.get("country") or e.get("setting"), 40),
            "Design (Sample)": (
                _clean_field(e.get("design"), 50)
                + (f"; n={_clean_field(e.get('sample'), 20)}" if e.get("sample") else "")
            ),
            "AI Tool": _clean_field(e.get("aiTool"), 40),
            "Educational Outcome": _clean_field(e.get("educationalOutcome"), 70),
            "Key Finding": _clean_field(e.get("keyFinding"), 100),
            "Limitation": _clean_field(e.get("mainLimitation"), 70),
            "Kirkpatrick": c.get("kp", "N/A"),
            "PMID/DOI": c.get("pmid", ""),
        })
    return pd.DataFrame(rows)


def build_overview_table_md(corpus, config):
    """Build a markdown overview table grouped by tier (manuscript style)."""
    if not corpus:
        return ""

    tier_order = {t: i for i, t in enumerate(config["tiers"])}
    grouped = {}
    for c in corpus:
        tier = c.get("tier", "") or "(Untiered)"
        grouped.setdefault(tier, []).append(c)

    # Order tiers per config, untiered last
    ordered_tiers = list(config["tiers"]) + [t for t in grouped if t not in config["tiers"]]

    md = []
    counter = 1
    for tier in ordered_tiers:
        studies = grouped.get(tier, [])
        if not studies:
            continue
        studies_sorted = sorted(studies, key=lambda c: c.get("year", ""))
        md.append(f"\n### {tier} (n = {len(studies)})\n")
        md.append("| # | Study | Country | Design (Sample) | AI Tool | Educational Outcome | Key Finding | Kirkpatrick |")
        md.append("|---|---|---|---|---|---|---|---|")
        for c in studies_sorted:
            e = c.get("extraction") or {}
            authors_short = _short_authors(c.get("authors", ""))
            study_cell = f"{authors_short} ({c.get('year','')})"
            country = _clean_field(e.get("country") or e.get("setting"), 30).replace("|", "/")
            design = _clean_field(e.get("design"), 40).replace("|", "/")
            sample = _clean_field(e.get("sample"), 20).replace("|", "/")
            design_cell = f"{design}" + (f"; n={sample}" if sample != "—" else "")
            ai_tool = _clean_field(e.get("aiTool"), 30).replace("|", "/")
            outcome = _clean_field(e.get("educationalOutcome"), 60).replace("|", "/")
            finding = _clean_field(e.get("keyFinding"), 80).replace("|", "/")
            kp = c.get("kp", "N/A")
            md.append(f"| {counter} | {study_cell} | {country} | {design_cell} | {ai_tool} | {outcome} | {finding} | {kp} |")
            counter += 1
    return "\n".join(md)


def categorise_exclusion(rationale):
    """Categorise an exclusion rationale into a standard reason group.
    Returns a short reason label suitable for the PRISMA flowchart."""
    if not rationale:
        return "Other / unspecified"
    r = rationale.lower()
    # Order matters — more specific patterns first
    if any(t in r for t in ["clinical accuracy", "diagnostic accuracy", "ai accuracy"]):
        return "Clinical AI accuracy (no educational outcome)"
    if any(t in r for t in ["llm benchmark", "mcq", "multiple choice", "benchmark"]):
        return "LLM/MCQ benchmark (no curricular implications)"
    if any(t in r for t in ["editorial", "commentary", "letter", "non-peer", "opinion"]):
        return "Editorial / commentary / non-peer-reviewed"
    if any(t in r for t in ["pilot", "n < 30", "small sample", "single class"]):
        return "Small pilot (no institutional scope)"
    if any(t in r for t in ["kap", "knowledge attitude practice", "without curricular"]):
        return "KAP without curricular recommendations"
    if any(t in r for t in ["tangential", "not substantive", "mention only", "passing"]):
        return "Tangential AI mention"
    if any(t in r for t in ["language", "non-english"]):
        return "Non-English / language barrier"
    if any(t in r for t in ["duplicate", "already included"]):
        return "Duplicate record"
    if "criterion" in r or "criteria" in r:
        return "Did not meet inclusion criteria"
    return "Other / unspecified"


def compute_prisma_numbers(results, corpus):
    """Compute PRISMA-ScR numbers from staged screening results and final corpus.
    Tracks Title → Abstract → Full-text + Appraisal stages."""
    if not results:
        return None

    # Identification — by database
    by_db = {}
    for r in results:
        by_db[r["db"]] = by_db.get(r["db"], 0) + 1

    n_after_dedup = len(results)

    # ─── Stage 1: Title screening ───
    n_s1_screened = sum(1 for r in results if r.get("title_decision", "Pending") != "Pending")
    n_s1_excluded = sum(1 for r in results if r.get("title_decision") == "Exclude")
    n_s1_included = sum(1 for r in results if r.get("title_decision") == "Include")
    n_s1_maybe = sum(1 for r in results if r.get("title_decision") == "Maybe")

    # ─── Stage 2: Abstract screening ───
    n_s2_screened = sum(1 for r in results
                        if r.get("title_decision") == "Include"
                        and r.get("abstract_decision", "Pending") != "Pending")
    n_s2_excluded = sum(1 for r in results if r.get("abstract_decision") == "Exclude")
    n_s2_included = sum(1 for r in results if r.get("abstract_decision") == "Include")
    n_s2_maybe = sum(1 for r in results if r.get("abstract_decision") == "Maybe")

    # ─── Stage 3: Full-text + Critical Appraisal ───
    n_s3_screened = sum(1 for r in results
                        if r.get("abstract_decision") == "Include"
                        and r.get("fulltext_decision", "Pending") != "Pending")
    n_s3_excluded = sum(1 for r in results if r.get("fulltext_decision") == "Exclude")
    n_s3_included = sum(1 for r in results if r.get("fulltext_decision") == "Include")

    # Exclusion reason breakdown — aggregate across stages
    exclusion_reasons_s1 = {}
    exclusion_reasons_s2 = {}
    exclusion_reasons_s3 = {}
    for r in results:
        if r.get("title_decision") == "Exclude":
            reason = categorise_exclusion(r.get("rationale", ""))
            exclusion_reasons_s1[reason] = exclusion_reasons_s1.get(reason, 0) + 1
        if r.get("abstract_decision") == "Exclude":
            reason = categorise_exclusion(r.get("rationale", ""))
            exclusion_reasons_s2[reason] = exclusion_reasons_s2.get(reason, 0) + 1
        if r.get("fulltext_decision") == "Exclude":
            reason = categorise_exclusion(r.get("rationale", ""))
            exclusion_reasons_s3[reason] = exclusion_reasons_s3.get(reason, 0) + 1

    # Fallback: legacy single-stage (if stages weren't used)
    n_pending = sum(1 for r in results if r.get("decision", "Pending") == "Pending")
    n_excluded_legacy = sum(1 for r in results if r.get("decision") == "Exclude")
    n_included_legacy = sum(1 for r in results if r.get("decision") == "Include")
    n_maybe_legacy = sum(1 for r in results if r.get("decision") == "Maybe")
    exclusion_reasons_legacy = {}
    for r in results:
        if r.get("decision") == "Exclude":
            reason = categorise_exclusion(r.get("rationale", ""))
            exclusion_reasons_legacy[reason] = exclusion_reasons_legacy.get(reason, 0) + 1

    n_corpus = len(corpus)

    return {
        "by_db": by_db,
        "n_after_dedup": n_after_dedup,
        # Stage 1
        "n_s1_screened": n_s1_screened,
        "n_s1_excluded": n_s1_excluded,
        "n_s1_included": n_s1_included,
        "n_s1_maybe": n_s1_maybe,
        "exclusion_reasons_s1": exclusion_reasons_s1,
        # Stage 2
        "n_s2_screened": n_s2_screened,
        "n_s2_excluded": n_s2_excluded,
        "n_s2_included": n_s2_included,
        "n_s2_maybe": n_s2_maybe,
        "exclusion_reasons_s2": exclusion_reasons_s2,
        # Stage 3
        "n_s3_screened": n_s3_screened,
        "n_s3_excluded": n_s3_excluded,
        "n_s3_included": n_s3_included,
        "exclusion_reasons_s3": exclusion_reasons_s3,
        # Final
        "n_corpus": n_corpus,
        # Legacy compatibility for older code paths
        "n_screened": n_s1_screened or (n_after_dedup - n_pending),
        "n_pending": n_pending,
        "n_excluded_screening": n_excluded_legacy if n_s1_excluded == 0 else (n_s1_excluded + n_s2_excluded + n_s3_excluded),
        "n_maybe": n_s1_maybe + n_s2_maybe if n_s1_maybe + n_s2_maybe > 0 else n_maybe_legacy,
        "n_included_screening": n_included_legacy if n_s3_included == 0 else n_s3_included,
        "exclusion_reasons": exclusion_reasons_legacy if not any([exclusion_reasons_s1, exclusion_reasons_s2, exclusion_reasons_s3])
                             else {k: exclusion_reasons_s1.get(k, 0) + exclusion_reasons_s2.get(k, 0) + exclusion_reasons_s3.get(k, 0)
                                   for k in set(list(exclusion_reasons_s1) + list(exclusion_reasons_s2) + list(exclusion_reasons_s3))},
    }


def _format_reasons(reasons_dict, max_show=4):
    """Format exclusion reasons as a bulleted block for Mermaid node text."""
    if not reasons_dict:
        return "(no exclusions yet)"
    sorted_r = sorted(reasons_dict.items(), key=lambda x: -x[1])
    lines = []
    for reason, count in sorted_r[:max_show]:
        safe = reason.replace("|", "/").replace('"', "'")[:50]
        lines.append(f"• {safe}: n={count}")
    if len(sorted_r) > max_show:
        other = sum(c for _, c in sorted_r[max_show:])
        lines.append(f"• Other: n={other}")
    return "<br/>".join(lines)


def build_prisma_flowchart(prisma):
    """Build a PRISMA-ScR style Mermaid flowchart with all 3 screening stages."""
    if not prisma:
        return "flowchart TD\n    A[No data yet]"

    # Identification block — one node per database
    db_lines = []
    for i, (db, n) in enumerate(sorted(prisma["by_db"].items(), key=lambda x: -x[1])):
        db_lines.append(f'    DB{i}["{db}<br/>n = {n:,}"]:::ident')
    db_to_dedup = "\n".join([f"    DB{i} --> Dedup" for i in range(len(prisma["by_db"]))])

    # Determine if we're using staged or legacy mode
    used_stages = prisma["n_s1_screened"] > 0 or prisma["n_s2_screened"] > 0 or prisma["n_s3_screened"] > 0

    if used_stages:
        # Build 3-stage flowchart
        excl_s1 = _format_reasons(prisma["exclusion_reasons_s1"])
        excl_s2 = _format_reasons(prisma["exclusion_reasons_s2"])
        excl_s3 = _format_reasons(prisma["exclusion_reasons_s3"])

        chart = f"""flowchart TD
    subgraph IDENT [" Identification "]
{chr(10).join(db_lines)}
    end

    Dedup["Records after duplicates removed<br/>n = {prisma['n_after_dedup']:,}"]:::stage

    S1Screened["Stage 1: Title screening<br/>n screened = {prisma['n_s1_screened']:,}"]:::stage
    S1Excluded["Excluded at title stage<br/>n = {prisma['n_s1_excluded']:,}<br/><br/>Reasons:<br/>{excl_s1}"]:::excluded
    S1Maybe["Stage 1 Maybe<br/>n = {prisma['n_s1_maybe']:,}<br/>(reviewed at Stage 2)"]:::maybe

    S2Screened["Stage 2: Abstract + journal quality<br/>n screened = {prisma['n_s2_screened']:,}"]:::stage
    S2Excluded["Excluded at abstract stage<br/>n = {prisma['n_s2_excluded']:,}<br/><br/>Reasons:<br/>{excl_s2}"]:::excluded
    S2Maybe["Stage 2 Maybe<br/>n = {prisma['n_s2_maybe']:,}"]:::maybe

    S3Screened["Stage 3: Full-text + critical appraisal<br/>n assessed = {prisma['n_s3_screened']:,}"]:::stage
    S3Excluded["Excluded at full-text stage<br/>n = {prisma['n_s3_excluded']:,}<br/><br/>Reasons (incl. appraisal):<br/>{excl_s3}"]:::excluded

    Included["Studies included in synthesis<br/>n = {prisma['n_corpus']:,}"]:::included

{db_to_dedup}
    Dedup --> S1Screened
    S1Screened --> S1Excluded
    S1Screened --> S1Maybe
    S1Screened --> S2Screened
    S1Maybe --> S2Screened
    S2Screened --> S2Excluded
    S2Screened --> S2Maybe
    S2Screened --> S3Screened
    S3Screened --> S3Excluded
    S3Screened --> Included

    classDef ident fill:#E6F1FB,stroke:#185FA5,color:#0C447C
    classDef stage fill:#F1EFE8,stroke:#5F5E5A,color:#2C2C2A
    classDef excluded fill:#FAECE7,stroke:#993C1D,color:#712B13
    classDef maybe fill:#FFF8DC,stroke:#8B7500,color:#5C4D00
    classDef included fill:#E1F5EE,stroke:#0F6E56,color:#085041
"""
    else:
        # Legacy single-stage flowchart (backward compat)
        excl = _format_reasons(prisma["exclusion_reasons"])
        chart = f"""flowchart TD
    subgraph IDENT [" Identification "]
{chr(10).join(db_lines)}
    end

    Dedup["Records after duplicates removed<br/>n = {prisma['n_after_dedup']:,}"]:::stage

    Screened["Records screened<br/>n = {prisma['n_screened']:,}"]:::stage

    Excluded["Records excluded at screening<br/>n = {prisma['n_excluded_screening']:,}<br/><br/>Reasons:<br/>{excl}"]:::excluded

    Eligible["Records assessed for eligibility<br/>n = {prisma['n_included_screening'] + prisma['n_maybe']:,}"]:::stage

    Uncertain["Records uncertain / Maybe<br/>n = {prisma['n_maybe']:,}"]:::excluded

    Included["Studies included in synthesis<br/>n = {prisma['n_corpus']:,}"]:::included

{db_to_dedup}
    Dedup --> Screened
    Screened --> Excluded
    Screened --> Eligible
    Eligible --> Uncertain
    Eligible --> Included

    classDef ident fill:#E6F1FB,stroke:#185FA5,color:#0C447C
    classDef stage fill:#F1EFE8,stroke:#5F5E5A,color:#2C2C2A
    classDef excluded fill:#FAECE7,stroke:#993C1D,color:#712B13
    classDef included fill:#E1F5EE,stroke:#0F6E56,color:#085041
"""
    return chart


def build_mermaid_flowchart(stats, config):
    """Legacy conceptual flowchart — kept for backward compatibility."""
    topic = config["research_topic"][:60]
    top_tiers = sorted(
        [(t, c) for t, c in stats["tier_counts"].items() if c > 0],
        key=lambda x: -x[1]
    )[:5]
    tier_nodes = "\n".join([f'    T{i}["{t} (n={c})"]' for i, (t, c) in enumerate(top_tiers)])
    tier_links = "\n".join([f"    Anchor --> T{i}" for i in range(len(top_tiers))])
    has_l3 = stats["kp_counts"].get("L3 Behaviour", 0) > 0
    has_l4 = stats["kp_counts"].get("L4 Results", 0) > 0
    if has_l4:
        ceiling = "L4 Results reached"
    elif has_l3:
        ceiling = "L3 Behaviour reached"
    else:
        ceiling = "L1/L2 only"
    tier_to_imp = "\n".join([f"    T{i} --> Ceiling" for i in range(len(top_tiers))])
    return f"""flowchart TD
    Anchor["Theoretical anchor:<br/>{topic}"]
{tier_nodes}
    Ceiling["Evidence ceiling:<br/>{ceiling}"]
    Policy["Policy implications"]
{tier_links}
{tier_to_imp}
    Ceiling --> Policy"""


def render_mermaid(diagram_code, height=420):
    """Render a Mermaid diagram inside Streamlit via embedded HTML."""
    html = f"""
    <div style="background:white;border-radius:8px;padding:1rem;">
      <pre class="mermaid" style="background:transparent;">
{diagram_code}
      </pre>
    </div>
    <script src="https://cdn.jsdelivr.net/npm/mermaid@10/dist/mermaid.min.js"></script>
    <script>
      mermaid.initialize({{ startOnLoad: true, theme: 'default', securityLevel: 'loose' }});
    </script>
    """
    st.components.v1.html(html, height=height, scrolling=False)


def synthesise_corpus(corpus, config, themes=None):
    """Generate structured synthesis: pre-computed tables + AI narrative in user's voice.
    If themes (list of dicts) is provided, the synthesis is organized by emergent themes
    instead of pre-defined tiers."""
    if not corpus:
        return None, "Corpus is empty."

    stats = compute_corpus_stats(corpus, config)
    use_themes = bool(themes)

    # Build study listing for narrative context
    listing = "\n".join([
        f"- {c['authors']} ({c['year']}) [{c.get('tier','?')}] [{c.get('kp','N/A')}] [{c['db']}]: {c['title']}"
        + (f" | Country: {c['extraction'].get('country','')}" if c.get("extraction") and c["extraction"].get("country") else "")
        + (f" | Finding: {c['extraction']['keyFinding']}" if c.get("extraction") and c["extraction"].get("keyFinding") else "")
        for c in corpus
    ])

    # Build organizing-structure block
    if use_themes:
        organizing_block = "EMERGENT THEMES (use these to structure the synthesis):\n" + "\n".join([
            f"  Theme {i+1} — {t.get('name','')}\n"
            f"    Definition: {t.get('definition','')}\n"
            f"    Supporting studies: {', '.join(t.get('supporting_studies', [])[:8])}\n"
            f"    Theoretical link: {t.get('theoretical_link','')}\n"
            for i, t in enumerate(themes)
        ])
    else:
        organizing_block = f"""Tier distribution:
{chr(10).join([f"  • {t}: {c}" for t, c in stats['tier_counts'].items() if c > 0])}

Kirkpatrick level distribution:
{chr(10).join([f"  • {k}: {c}" for k, c in stats['kp_counts'].items() if c > 0])}"""

    # Pass pre-computed stats to Claude as context
    stats_summary = f"""PRE-COMPUTED CORPUS STATISTICS (use these exact numbers in your narrative — do not recompute):

Total studies: {stats['n']}
Year range: {stats['year_range']}

{organizing_block}

Database sourcing:
{chr(10).join([f"  • {d}: {c}" for d, c in sorted(stats['db_counts'].items(), key=lambda x: -x[1])])}

Geographic distribution (from extractions):
{chr(10).join([f"  • {g}: {c}" for g, c in sorted(stats['geo_counts'].items(), key=lambda x: -x[1])[:10]]) if stats['geo_counts'] else "  • (No extraction data — geographic synthesis will be limited)"}

Study design distribution (from extractions):
{chr(10).join([f"  • {d}: {c}" for d, c in sorted(stats['design_counts'].items(), key=lambda x: -x[1])[:8]]) if stats['design_counts'] else "  • (No extraction data)"}
"""

    user_msg = stats_summary + "\n\nCORPUS LISTING:\n" + listing
    system_prompt = build_synth_system_prompt(use_themes=use_themes, themes=themes)
    text, err = call_claude(system_prompt, user_msg, max_tokens=6000, use_cache=True)

    if err:
        return None, err

    return {"narrative": text, "stats": stats, "themes": themes, "use_themes": use_themes}, None


# ─────────────────────────────────────────────────────────────────
# THEME DISCOVERY (emergent themes from extracted data)
# ─────────────────────────────────────────────────────────────────

def build_theme_discovery_prompt():
    c = st.session_state.config
    return f"""You are a qualitative analyst conducting an inductive thematic analysis on extracted study data for a scoping review on: {c['research_topic']}.

Theoretical anchor: {c['theoretical_anchor']}

TASK
You will receive the extracted findings from every included study. Conduct an inductive thematic analysis — themes must emerge from the data, NOT be imposed from a pre-existing framework. Identify 4-7 distinct themes that capture the most salient patterns across the corpus.

GUIDELINES
- Each theme should be data-driven and supported by 3 or more studies
- Themes should be conceptually distinct (minimal overlap)
- Theme names should be concise and analytically descriptive (5-9 words), not generic ("Findings", "Results")
- For each theme, identify which studies support it (by author + year)
- Tentatively note how each theme relates to the theoretical anchor
- Avoid assertive language — describe what the studies "appear to suggest" or "tentatively indicate"

Respond ONLY as valid JSON, no markdown, no preamble:
{{
  "themes": [
    {{
      "name": "Concise analytical theme name",
      "definition": "One-to-two-sentence definition framed tentatively (e.g., 'The findings appear to indicate that...')",
      "supporting_studies": ["Author Year", "Author Year", "Author Year"],
      "theoretical_link": "One sentence on how this theme tentatively relates to the theoretical anchor",
      "tensions": "One sentence noting any contradictions or open questions within this theme (or empty string if none)"
    }}
  ]
}}"""


def discover_themes(corpus):
    """Run inductive theme discovery across the extracted corpus."""
    if not corpus:
        return None, "Corpus is empty — add studies first."

    extracted = [c for c in corpus if c.get("extraction")]
    if len(extracted) < 3:
        return None, f"Theme discovery needs at least 3 extracted studies (have {len(extracted)}). Run AI extraction first."

    # Build compact study summaries for the analysis
    lines = []
    for c in extracted:
        e = c["extraction"]
        author = (c.get("authors") or "").split(",")[0].strip()
        year = c.get("year", "")
        finding = e.get("keyFinding") or ""
        outcome = e.get("educationalOutcome") or ""
        country = e.get("country") or ""
        design = e.get("design") or ""
        lines.append(
            f"- {author} ({year}) [{country}, {design}]: outcome = {outcome}; finding = {finding}"
        )

    user_msg = f"EXTRACTED CORPUS (n={len(extracted)} studies):\n\n" + "\n".join(lines)
    text, err = call_claude(build_theme_discovery_prompt(), user_msg, max_tokens=3000)

    if err:
        return None, err

    try:
        cleaned = text.replace("```json", "").replace("```", "").strip()
        data = json.loads(cleaned)
        themes = data.get("themes", [])
        if not themes:
            return None, "AI returned no themes — try running again with more extracted studies."
        return themes, None
    except Exception as e:
        return None, f"Could not parse theme JSON: {e}\n\nRaw response (first 500 chars): {text[:500]}"


# ─────────────────────────────────────────────────────────────────
# CONVERSATIONAL REFINEMENT (chat with the synthesis)
# ─────────────────────────────────────────────────────────────────

def build_refinement_prompt():
    c = st.session_state.config
    style_section = ""
    if c.get("writing_style_sample", "").strip():
        style_section = f"""

The author's writing style sample (mimic precisely if rewriting prose):
\"\"\"
{c['writing_style_sample'][:1500]}
\"\"\"
"""

    return f"""You are a synthesis collaborator helping refine a scoping review on: {c['research_topic']}.

Theoretical anchor: {c['theoretical_anchor']}
{style_section}

You will receive the current synthesis text and a user request. Respond conversationally — directly address what the user asked. If they ask you to rewrite a section, return the rewritten section in proper prose (no headers needed). If they ask a question, answer it. If they want you to add a counter-argument, draft it. If they want less hedging, rewrite with measured (but not absolute) language.

HARD RULES
- Always remain in measured academic voice. Avoid "proves", "clearly", "definitively", "must", "always", "never", "the evidence shows" (use "the evidence suggests").
- Mimic the author's writing style (sentence rhythm, hedging patterns, connective adverbs Moreover/Therefore/However/Interestingly).
- Do not fabricate studies or findings beyond what is in the corpus listing.
- If the user asks for something the corpus cannot support, say so honestly and offer the closest defensible alternative.
- Keep responses focused — answer the specific question or rewrite the specific section requested. Do not regenerate the whole synthesis unless explicitly asked."""


def refine_synthesis(current_narrative, chat_history, user_message, corpus):
    """Refine the synthesis based on a user instruction.
    Returns (response_text, error)."""

    # Build compact corpus listing for context
    corpus_lines = []
    for c in corpus[:50]:  # cap for context window
        author = (c.get("authors") or "").split(",")[0].strip()
        year = c.get("year", "")
        e = c.get("extraction") or {}
        finding = (e.get("keyFinding") or c.get("title") or "")[:120]
        corpus_lines.append(f"- {author} ({year}): {finding}")

    # Build conversation context
    history_text = ""
    for turn in chat_history[-6:]:  # last 6 turns to manage context
        history_text += f"\n[USER]: {turn['user']}\n[ASSISTANT]: {turn['assistant'][:800]}\n"

    user_msg = f"""CURRENT SYNTHESIS:
\"\"\"
{current_narrative[:8000]}
\"\"\"

CORPUS (for grounding — n={len(corpus)}):
{chr(10).join(corpus_lines)}

PRIOR CONVERSATION:
{history_text if history_text else "(none yet)"}

USER REQUEST:
{user_message}

Respond directly to the request. If rewriting a section, return the rewritten prose. If answering a question, give a focused answer."""

    text, err = call_claude(build_refinement_prompt(), user_msg, max_tokens=3000)
    if err:
        return None, err
    return text, None


# ─────────────────────────────────────────────────────────────────
# JOURNAL TRENDS DISCOVERY
# ─────────────────────────────────────────────────────────────────

DENTAL_JOURNAL_SHORTLIST = [
    "Journal of Dental Education",
    "European Journal of Dental Education",
    "Journal of Dentistry",
    "Journal of Dental Research",
    "International Dental Journal",
    "BMC Medical Education",
    "BMC Oral Health",
    "JMIR Medical Education",
    "Medical Education",
    "Journal of the American Dental Association",
    "British Dental Journal",
    "Caries Research",
    "Clinical Oral Investigations",
    "Frontiers in Dental Medicine",
]


def search_journals_recent(journals, years_back, max_per_journal):
    """Pull recent papers from a list of journals using PubMed [Journal] tag.
    Returns combined deduplicated list of papers from each journal."""
    from datetime import datetime as _dt
    current_year = _dt.now().year
    date_from = current_year - years_back
    date_to = current_year

    all_items = []
    per_journal_counts = {}
    for j in journals:
        # Use PubMed Journal tag — exact-match by title
        # Use sort order = pub date desc via the API
        query = f'"{j}"[Journal]'
        items, total = search_pubmed(query, date_from, date_to, max_per_journal)
        per_journal_counts[j] = {"retrieved": len(items), "total": total}
        all_items.extend(items)

    # Dedup by PMID
    seen = set()
    unique = []
    for it in all_items:
        pmid = it.get("pmid") or it.get("id")
        if pmid and pmid not in seen:
            seen.add(pmid)
            unique.append(it)

    return unique, per_journal_counts


def build_trends_prompt():
    c = st.session_state.config
    return f"""You are a research-trends analyst examining a recent corpus of papers from major journals in a specific field.

Research focus: {c['research_topic']}

TASK
You will receive a list of titles + (where available) abstracts from recent journal publications. Identify 5-8 distinct emerging trends or themes that appear in this corpus. For each trend:
- Give a concise analytical name (5-9 words)
- Provide a one-sentence definition framed tentatively
- Cite 3-5 supporting paper titles or first-author-year tags
- Note tentative significance for the research focus

Avoid assertive phrasing. Use hedging ("appears to", "tentatively", "may indicate").

Respond ONLY as valid JSON:
{{
  "trends": [
    {{
      "name": "Concise trend name",
      "definition": "Tentative one-sentence definition",
      "supporting_papers": ["First-author Year", "First-author Year"],
      "significance": "One sentence on why this trend may matter for the research focus"
    }}
  ],
  "summary": "Two-sentence overall summary of the trends landscape"
}}"""


def analyze_journal_trends(papers):
    """Use Claude to extract emerging trends from a corpus of recent journal papers."""
    if not papers:
        return None, "No papers to analyze."

    lines = []
    for p in papers[:200]:  # cap for context
        author = (p.get("authors") or "").split(",")[0].strip()
        year = p.get("year", "")
        title = p.get("title", "")
        abstract = (p.get("abstract") or "")[:300]
        lines.append(f"- {author} ({year}) [{p.get('journal','')}]: {title}")
        if abstract:
            lines.append(f"    Abstract: {abstract}")

    user_msg = f"CORPUS OF RECENT JOURNAL PAPERS (n={len(papers)}):\n\n" + "\n".join(lines)
    text, err = call_claude(build_trends_prompt(), user_msg, max_tokens=4000)
    if err:
        return None, err

    try:
        cleaned = text.replace("```json", "").replace("```", "").strip()
        data = json.loads(cleaned)
        return data, None
    except Exception as e:
        return None, f"Could not parse trends JSON: {e}\n\nRaw (first 500 chars): {text[:500]}"


# ─────────────────────────────────────────────────────────────────
# WORD DOCUMENT EXPORT
# ─────────────────────────────────────────────────────────────────

def build_word_document(corpus, stats, narrative, prisma, config):
    """Build a .docx document with full synthesis content.
    Returns BytesIO object suitable for Streamlit download_button.
    """
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from io import BytesIO

    doc = Document()

    # ─── Set default font ───
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ─── Title ───
    title = doc.add_heading(f"Synthesis: {config['research_topic']}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(f"Generated {datetime.now().strftime('%Y-%m-%d %H:%M')}  ·  n = {stats['n']} studies  ·  Year range {stats['year_range']}")
    meta_run.italic = True
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()  # spacer

    # ─── PRISMA-ScR Numbers ───
    if prisma:
        doc.add_heading("PRISMA-ScR Flow Summary", level=1)
        p = doc.add_paragraph()
        p.add_run("Records identified across databases: ").bold = True
        p.add_run(", ".join([f"{db} (n = {n:,})" for db, n in sorted(prisma['by_db'].items(), key=lambda x: -x[1])]))

        p = doc.add_paragraph()
        p.add_run("Records after duplicates removed: ").bold = True
        p.add_run(f"n = {prisma['n_after_dedup']:,}")

        p = doc.add_paragraph()
        p.add_run("Records screened: ").bold = True
        p.add_run(f"n = {prisma['n_screened']:,}")

        p = doc.add_paragraph()
        p.add_run("Records excluded at screening: ").bold = True
        p.add_run(f"n = {prisma['n_excluded_screening']:,}")

        if prisma['exclusion_reasons']:
            doc.add_paragraph("Exclusion reasons:", style='Intense Quote')
            for reason, count in sorted(prisma['exclusion_reasons'].items(), key=lambda x: -x[1]):
                bullet = doc.add_paragraph(style='List Bullet')
                bullet.add_run(f"{reason}: n = {count}")

        p = doc.add_paragraph()
        p.add_run("Records uncertain (Maybe / deferred): ").bold = True
        p.add_run(f"n = {prisma['n_maybe']:,}")

        p = doc.add_paragraph()
        p.add_run("Studies included in synthesis: ").bold = True
        run = p.add_run(f"n = {prisma['n_corpus']:,}")
        run.bold = True
        run.font.color.rgb = RGBColor(0x0F, 0x6E, 0x56)

        doc.add_paragraph()

    # ─── Overview Table (Table 1) ───
    doc.add_heading("Table 1. Overview of Included Studies", level=1)
    overview_df = build_overview_table_df(corpus, config)
    if overview_df is not None and not overview_df.empty:
        # Use a compact column subset for Word to avoid overflow
        word_cols = ["#", "Tier", "Study", "Country/Setting", "Design (Sample)", "AI Tool", "Educational Outcome", "Key Finding", "Kirkpatrick"]
        cols = [c for c in word_cols if c in overview_df.columns]
        table = doc.add_table(rows=1, cols=len(cols))
        table.style = 'Light Grid Accent 1'
        table.autofit = True

        # Header row
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(cols):
            cell = hdr_cells[i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(col)
            run.bold = True
            run.font.size = Pt(9)

        # Body rows
        for _, row in overview_df.iterrows():
            cells = table.add_row().cells
            for i, col in enumerate(cols):
                val = str(row[col]) if pd.notna(row[col]) else "—"
                cells[i].text = ""
                run = cells[i].paragraphs[0].add_run(val)
                run.font.size = Pt(8)

        # Tighten column widths approximately
        widths_cm = {
            "#": 0.6, "Tier": 1.8, "Study": 2.2, "Country/Setting": 1.8,
            "Design (Sample)": 2.5, "AI Tool": 1.8, "Educational Outcome": 2.5,
            "Key Finding": 3.5, "Kirkpatrick": 1.4,
        }
        for i, col in enumerate(cols):
            w = widths_cm.get(col, 2.0)
            for cell in table.columns[i].cells:
                cell.width = Cm(w)

    doc.add_page_break()

    # ─── Statistical Tables Section ───
    doc.add_heading("Statistical Tables", level=1)

    # Tier distribution table
    doc.add_heading("Table 2. Tier Distribution", level=2)
    tier_table = doc.add_table(rows=1, cols=3)
    tier_table.style = 'Light Grid Accent 1'
    hdr = tier_table.rows[0].cells
    hdr[0].text = "Tier"
    hdr[1].text = "n"
    hdr[2].text = "% of corpus"
    for cell in hdr:
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for tier, count in stats["tier_counts"].items():
        if count == 0:
            continue
        pct = round(100 * count / stats["n"], 1) if stats["n"] else 0
        row = tier_table.add_row().cells
        row[0].text = tier
        row[1].text = str(count)
        row[2].text = f"{pct}%"

    doc.add_paragraph()

    # Kirkpatrick distribution
    doc.add_heading("Table 3. Kirkpatrick Level Distribution", level=2)
    kp_table = doc.add_table(rows=1, cols=3)
    kp_table.style = 'Light Grid Accent 1'
    hdr = kp_table.rows[0].cells
    hdr[0].text = "Kirkpatrick level"
    hdr[1].text = "n"
    hdr[2].text = "% of corpus"
    for cell in hdr:
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for kp, count in stats["kp_counts"].items():
        if count == 0:
            continue
        pct = round(100 * count / stats["n"], 1) if stats["n"] else 0
        row = kp_table.add_row().cells
        row[0].text = kp
        row[1].text = str(count)
        row[2].text = f"{pct}%"

    doc.add_paragraph()

    # Geographic distribution if available
    if stats.get("geo_counts"):
        doc.add_heading("Table 4. Geographic Distribution", level=2)
        geo_table = doc.add_table(rows=1, cols=2)
        geo_table.style = 'Light Grid Accent 1'
        hdr = geo_table.rows[0].cells
        hdr[0].text = "Country / setting"
        hdr[1].text = "n"
        for cell in hdr:
            for run in cell.paragraphs[0].runs:
                run.bold = True
        for country, count in sorted(stats["geo_counts"].items(), key=lambda x: -x[1])[:15]:
            row = geo_table.add_row().cells
            row[0].text = country
            row[1].text = str(count)
        doc.add_paragraph()

    doc.add_page_break()

    # ─── Narrative Sections ───
    sections = [s.strip() for s in narrative.split("===SECTION===")]
    while len(sections) < 9:
        sections.append("(Section narrative not generated.)")

    section_titles = [
        ("Results", None),
        (None, "1. Corpus Overview"),
        (None, "2. Tier Distribution and Evidence Weight"),
        (None, "3. Kirkpatrick Outcome Ceiling"),
        (None, "4. Geographic and Contextual Variation"),
        (None, "5. Cross-Tier Patterns"),
        ("Discussion", None),
        (None, "6. Comparison with Available Literature"),
        (None, "7. Integration Across Tiers"),
        (None, "8. Argument and Implications"),
        ("Conclusion", None),
        (None, "9. Conclusion"),
    ]

    section_idx = 0
    for parent, sub in section_titles:
        if parent:
            doc.add_heading(parent, level=1)
        if sub:
            doc.add_heading(sub, level=2)
            if section_idx < len(sections):
                # Split narrative into paragraphs
                content = sections[section_idx]
                for para in content.split("\n\n"):
                    para = para.strip()
                    if para:
                        p = doc.add_paragraph(para)
                        p.paragraph_format.space_after = Pt(8)
                        for run in p.runs:
                            run.font.size = Pt(11)
                section_idx += 1

    # ─── Critical Appraisal Summary (if appraisals have been run) ───
    appraisals = st.session_state.get("appraisals", {})
    corpus_appraisals = {c["id"]: appraisals[c["id"]] for c in corpus if c["id"] in appraisals}
    if corpus_appraisals:
        doc.add_page_break()
        doc.add_heading("Critical Appraisal Summary", level=1)
        # Summary table
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'
        hdr = table.rows[0].cells
        hdr[0].text = "Study"
        hdr[1].text = "Appraisal tool"
        hdr[2].text = "Overall rating"
        hdr[3].text = "Recommendation"
        for cell in hdr:
            for r in cell.paragraphs[0].runs:
                r.bold = True
        for c in corpus:
            ap = corpus_appraisals.get(c["id"])
            if ap:
                row = table.add_row().cells
                row[0].text = f"{(c.get('authors') or '').split(',')[0].strip()} ({c.get('year','')})"
                row[1].text = ap.get("tool", "")
                row[2].text = ap.get("overall_rating", "")
                row[3].text = ap.get("recommendation", "")

    # ─── References (Manchester-Harvard) ───
    refs = build_reference_list(corpus)
    if refs:
        doc.add_page_break()
        doc.add_heading("References", level=1)
        meta = doc.add_paragraph()
        meta_run = meta.add_run("Formatted in Manchester-Harvard style.")
        meta_run.italic = True
        meta_run.font.size = Pt(10)
        meta_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        for ref in refs:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.left_indent = Cm(1.0)
            p.paragraph_format.first_line_indent = Cm(-1.0)  # hanging indent
            # Render with italics for journal title (delimited by *...* in our format)
            import re as _re
            parts = _re.split(r"(\*[^*]+\*)", ref)
            for part in parts:
                if part.startswith("*") and part.endswith("*") and len(part) > 2:
                    run = p.add_run(part[1:-1])
                    run.italic = True
                else:
                    p.add_run(part)
            for run in p.runs:
                run.font.size = Pt(11)

    # Save to BytesIO
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


def call_claude(system_prompt, user_message, max_tokens=1000, model="claude-sonnet-4-5", use_cache=False):
    """Call the Claude API.

    use_cache=True adds cache_control to the system prompt block so repeated calls
    with the same system prompt only charge ~10% of system-prompt tokens after the
    first call (5-minute TTL). Use for any function that loops over many articles
    with an identical system prompt.
    """
    client = get_anthropic_client()
    if not client:
        return None, "No Anthropic API key configured. Add ANTHROPIC_API_KEY to Streamlit secrets or environment."
    try:
        if use_cache:
            system_block = [{"type": "text", "text": system_prompt, "cache_control": {"type": "ephemeral"}}]
        else:
            system_block = system_prompt
        msg = client.messages.create(
            model=model,
            max_tokens=max_tokens,
            system=system_block,
            messages=[{"role": "user", "content": user_message}],
        )
        return msg.content[0].text, None
    except Exception as e:
        return None, str(e)


def screen_article(article):
    user_msg = f"""Title: {article['title']}
Authors: {article['authors']}
Journal: {article['journal']} ({article['year']})
Database: {article['db']}"""
    if article.get("abstract"):
        user_msg += f"\nAbstract: {article['abstract'][:600]}"
    # Haiku is sufficient for binary screening; cache the system prompt across calls
    text, err = call_claude(
        build_screen_system_prompt(), user_msg,
        model="claude-haiku-4-5-20251001", use_cache=True,
    )
    if err:
        return None, err
    try:
        cleaned = text.replace("```json", "").replace("```", "").strip()
        return json.loads(cleaned), None
    except Exception as e:
        return None, f"Parse error: {e}\nRaw: {text[:200]}"


def build_title_screen_prompt():
    """Title-only screening prompt — fast triage, errs on the side of inclusion (Maybe)."""
    c = st.session_state.config
    # Deliberately short: theoretical anchor and long criteria are not needed
    # for a simple title triage that errs toward Maybe.
    return f"""You are screening titles for a scoping review on: {c['research_topic']}.

Stage 1 triage — decide based on title ONLY. Err toward Maybe when uncertain.
- Include: title clearly relevant
- Maybe: ambiguous or unclear from title alone (DEFAULT for uncertainty)
- Exclude: title clearly off-topic (wrong field/population/concept)

Respond ONLY as valid JSON:
{{"decision":"Include|Maybe|Exclude","rationale":"1-line reason","confidence":"high|medium|low"}}"""


def batch_screen_title(articles):
    """Screen up to 20 articles in a single API call (title-only, Stage 1).
    Returns list of result dicts parallel to the input list, or None on error."""
    if not articles:
        return [], None

    lines = []
    for i, a in enumerate(articles, 1):
        lines.append(f'{i}. Title: {a["title"]} | Journal: {a.get("journal","")} ({a.get("year","")})')

    user_msg = (
        "Screen each title below. Return a JSON array with one object per record in the same order.\n"
        "Each object: {\"decision\":\"Include|Maybe|Exclude\",\"rationale\":\"1-line\",\"confidence\":\"high|medium|low\"}\n\n"
        + "\n".join(lines)
        + "\n\nReturn ONLY a JSON array, no markdown."
    )
    text, err = call_claude(
        build_title_screen_prompt(), user_msg,
        max_tokens=80 * len(articles),
        model="claude-haiku-4-5-20251001", use_cache=True,
    )
    if err:
        return None, err
    try:
        cleaned = text.replace("```json", "").replace("```", "").strip()
        results = json.loads(cleaned)
        if isinstance(results, list) and len(results) == len(articles):
            return results, None
        return None, f"Expected {len(articles)} results, got {len(results) if isinstance(results, list) else 'non-list'}"
    except Exception as e:
        return None, f"Parse error: {e}\nRaw: {text[:300]}"


def batch_screen_abstract(articles):
    """Screen up to 10 articles in a single API call (title+abstract, Stage 2).
    Returns list of result dicts parallel to the input list, or None on error."""
    if not articles:
        return [], None

    lines = []
    for i, a in enumerate(articles, 1):
        abstract = (a.get("abstract") or "")[:400]
        lines.append(
            f'{i}. Title: {a["title"]} | Journal: {a.get("journal","")} ({a.get("year","")})'
            + (f'\n   Abstract: {abstract}' if abstract else "")
        )

    user_msg = (
        "Screen each record below for inclusion. Return a JSON array with one object per record in the same order.\n"
        f'Each object: {{"decision":"Include|Maybe|Exclude","tier":"<tier or empty>","kp":"<KP level>","rationale":"1-sentence","confidence":"high|medium|low"}}\n\n'
        + "\n\n".join(lines)
        + "\n\nReturn ONLY a JSON array, no markdown."
    )
    text, err = call_claude(
        build_screen_system_prompt(), user_msg,
        max_tokens=120 * len(articles),
        model="claude-haiku-4-5-20251001", use_cache=True,
    )
    if err:
        return None, err
    try:
        cleaned = text.replace("```json", "").replace("```", "").strip()
        results = json.loads(cleaned)
        if isinstance(results, list) and len(results) == len(articles):
            return results, None
        return None, f"Expected {len(articles)} results, got {len(results) if isinstance(results, list) else 'non-list'}"
    except Exception as e:
        return None, f"Parse error: {e}\nRaw: {text[:300]}"


def screen_article_title_only(article):
    """Stage 1: title-only screening. Fast, conservative — errs toward Maybe."""
    user_msg = f'Title: {article["title"]} | Journal: {article.get("journal","")} ({article.get("year","")})'
    text, err = call_claude(
        build_title_screen_prompt(), user_msg,
        max_tokens=150, model="claude-haiku-4-5-20251001", use_cache=True,
    )
    if err:
        return None, err
    try:
        cleaned = text.replace("```json", "").replace("```", "").strip()
        return json.loads(cleaned), None
    except Exception as e:
        return None, f"Parse error: {e}\nRaw: {text[:200]}"


def extract_article(article):
    user_msg = f"""Title: {article['title']}
Authors: {article['authors']}
Journal: {article['journal']} ({article['year']})
Database: {article['db']}
Tier: {article.get('tier','?')}
Kirkpatrick: {article.get('kp','N/A')}
Screening rationale: {article.get('rationale','')}"""
    if article.get("abstract"):
        user_msg += f"\nAbstract: {article['abstract'][:600]}"
    # Extraction needs more nuance — keep Sonnet but cache system prompt
    text, err = call_claude(build_extract_system_prompt(), user_msg, use_cache=True)
    if err:
        return None, err
    try:
        cleaned = text.replace("```json", "").replace("```", "").strip()
        return json.loads(cleaned), None
    except Exception as e:
        return None, f"Parse error: {e}\nRaw: {text[:200]}"


# ─────────────────────────────────────────────────────────────────
# UI HELPERS
# ─────────────────────────────────────────────────────────────────

DECISION_COLORS = {
    "Include": "#0F6E56",
    "Maybe": "#854F0B",
    "Exclude": "#993C1D",
    "Pending": "#888780",
}
DECISION_BG = {
    "Include": "#E1F5EE",
    "Maybe": "#FAEEDA",
    "Exclude": "#FAECE7",
    "Pending": "#F1EFE8",
}
DB_COLORS = {
    "PubMed": ("#E6F1FB", "#0C447C"),
    "Europe PMC": ("#EAF3DE", "#27500A"),
    "OpenAlex": ("#EEEDFE", "#3C3489"),
}


def badge(text, bg, color, font_size="11px"):
    return f'<span style="display:inline-block;background:{bg};color:{color};padding:2px 8px;border-radius:6px;font-size:{font_size};font-weight:500;margin-right:4px">{text}</span>'


def db_badge(db):
    bg, col = DB_COLORS.get(db, ("#F1EFE8", "#444441"))
    return badge(db, bg, col)


# ─────────────────────────────────────────────────────────────────
# HEADER & SIDEBAR
# ─────────────────────────────────────────────────────────────────

st.markdown("""
<style>
.stApp { font-family: -apple-system, BlinkMacSystemFont, sans-serif; }
.block-container { padding-top: 1.5rem; max-width: 1200px; }
</style>
""", unsafe_allow_html=True)

st.title("📚 Literature Research Tool")
st.caption("Multi-database search · AI-powered screening · structured extraction · synthesis")

# API key status
client = get_anthropic_client()
if not client:
    st.error("⚠️ No Anthropic API key found. Add `ANTHROPIC_API_KEY` to Streamlit secrets (`.streamlit/secrets.toml`) or environment variables.")
    st.info("AI screening, extraction and synthesis features require the API key. Database search works without it.")

# Sidebar: Config
with st.sidebar:
    st.header("⚙️ Configuration")
    with st.expander("Research topic & criteria", expanded=False):
        st.session_state.config["research_topic"] = st.text_input(
            "Research topic",
            value=st.session_state.config["research_topic"],
        )
        st.session_state.config["theoretical_anchor"] = st.text_area(
            "Theoretical anchor (used in screening & synthesis)",
            value=st.session_state.config["theoretical_anchor"],
            height=100,
        )
        st.session_state.config["inclusion_criteria"] = st.text_area(
            "Inclusion criteria",
            value=st.session_state.config["inclusion_criteria"],
            height=150,
        )
        st.session_state.config["exclusion_criteria"] = st.text_area(
            "Exclusion criteria",
            value=st.session_state.config["exclusion_criteria"],
            height=150,
        )

    with st.expander("Tiers & Kirkpatrick", expanded=False):
        tiers_text = st.text_area(
            "Tiers (one per line)",
            value="\n".join(st.session_state.config["tiers"]),
            height=120,
        )
        st.session_state.config["tiers"] = [t.strip() for t in tiers_text.split("\n") if t.strip()]
        st.session_state.config["tier_descriptions"] = st.text_area(
            "Tier descriptions",
            value=st.session_state.config["tier_descriptions"],
            height=120,
        )
        kp_text = st.text_area(
            "Kirkpatrick levels (one per line)",
            value="\n".join(st.session_state.config["kp_levels"]),
            height=120,
        )
        st.session_state.config["kp_levels"] = [k.strip() for k in kp_text.split("\n") if k.strip()]
        st.session_state.config["kp_description"] = st.text_area(
            "Kirkpatrick descriptions",
            value=st.session_state.config["kp_description"],
            height=120,
        )

    with st.expander("✍️ Writing style", expanded=False):
        st.caption("Paste 2-3 paragraphs of your own academic writing. Claude will mimic your sentence rhythm, vocabulary, hedging language, and paragraph openers in the synthesis narrative.")
        st.session_state.config["writing_style_sample"] = st.text_area(
            "Writing style sample (your own published work)",
            value=st.session_state.config.get("writing_style_sample", ""),
            height=250,
            help="Best results: paste 300-500 words from your most representative paper. The voice features are extracted automatically.",
        )
        if st.session_state.config.get("writing_style_sample", "").strip():
            words = len(st.session_state.config["writing_style_sample"].split())
            st.caption(f"✓ {words} words loaded · style will be applied to synthesis narrative")
        else:
            st.warning("No style sample set — synthesis will use a generic academic voice.")

    with st.expander("🔑 API keys (optional)", expanded=False):
        st.caption("Add credentials for paid databases. Keys are stored only in this session and never logged.")
        st.session_state.elsevier_key = st.text_input(
            "Elsevier API key",
            value=st.session_state.get("elsevier_key", ""),
            type="password",
            help="Get a free key at https://dev.elsevier.com. Required for Scopus. EMBASE additionally requires institutional subscription.",
        )
        if st.session_state.elsevier_key:
            st.caption(f"✓ Key loaded ({len(st.session_state.elsevier_key)} chars)")
        else:
            st.caption("Scopus and EMBASE will be skipped at search time.")

    st.divider()
    if st.button("Reset to defaults"):
        st.session_state.config = DEFAULT_CONFIG.copy()
        st.rerun()

    st.divider()
    st.markdown("**Session stats**")
    st.metric("Results", len(st.session_state.results))
    st.metric("Corpus", len(st.session_state.corpus))

# ─────────────────────────────────────────────────────────────────
# TABS
# ─────────────────────────────────────────────────────────────────

tab_search, tab_screen, tab_corpus, tab_synth, tab_trends = st.tabs(["🔍 Search", "✓ Screen", "📂 Corpus", "📝 Synthesis", "📈 Trends"])

# ─── SEARCH TAB ────────────────────────────────────────────────
with tab_search:
    st.subheader("Database search")

    # ─── AI query builder + presets ───
    builder_cols = st.columns([3, 1, 1])
    with builder_cols[0]:
        st.caption("🤖 Let AI generate an optimal query from your topic + inclusion criteria, OR pick a preset below, OR write your own.")
    with builder_cols[1]:
        if st.button("🤖 Build query with AI", type="secondary", use_container_width=True, disabled=not client):
            if client:
                with st.spinner("Generating optimal query..."):
                    q_data, err = ai_generate_query()
                    if q_data and q_data.get("query"):
                        st.session_state.pending_query = q_data["query"]
                        st.session_state.last_query_meta = q_data
                        st.rerun()
                    else:
                        st.error(f"Query builder error: {err}")
    with builder_cols[2]:
        if "last_query_meta" in st.session_state and st.session_state.get("last_query_meta"):
            meta = st.session_state.last_query_meta
            st.caption(f"_{meta.get('estimated_yield','?')} ({meta.get('estimated_count','?')})_")

    if "last_query_meta" in st.session_state and st.session_state.get("last_query_meta"):
        with st.expander("AI query rationale", expanded=False):
            meta = st.session_state.last_query_meta
            st.markdown(f"**Rationale:** {meta.get('rationale','—')}")
            st.markdown(f"**Estimated yield:** {meta.get('estimated_yield','?')} (~{meta.get('estimated_count','?')} papers)")

    preset_cols = st.columns(min(5, len(st.session_state.config["search_presets"])))
    for i, (label, query) in enumerate(st.session_state.config["search_presets"]):
        with preset_cols[i % len(preset_cols)]:
            if st.button(label, key=f"preset_{i}", use_container_width=True):
                st.session_state.pending_query = query
                st.rerun()

    default_q = st.session_state.get("pending_query", st.session_state.config["search_presets"][0][1])
    query = st.text_area("Search query", value=default_q, height=100, key="search_query_input")

    c1, c2, c3, c4 = st.columns([1, 1, 1, 1])
    with c1:
        date_from = st.number_input("From year", min_value=1990, max_value=2030, value=2018)
    with c2:
        date_to = st.number_input("To year", min_value=1990, max_value=2030, value=2026)
    with c3:
        max_per_db = st.number_input(
            "Max per DB",
            min_value=10,
            max_value=2000,
            value=200,
            step=50,
            help="How many results to fetch from EACH database. Set higher (500-1000) for comprehensive sweeps like PRISMA-ScR systematic searches. Fetching 1000+ may take 1-2 minutes due to batched abstract retrieval.",
        )
    with c4:
        st.write("")
        st.write("")
        search_btn = st.button("🔍 Search all databases", type="primary", use_container_width=True)

    # ─── Preview count (cheap, fast diagnostic) ───
    preview_btn = st.button(
        "👁️ Preview total counts (fast, no abstracts)",
        type="secondary",
        help="Quickly check how many papers your query matches in each database before doing a full fetch. Useful for tuning the query when yields look low.",
    )
    if preview_btn and query.strip():
        with st.spinner("Checking counts across databases..."):
            counts = {}
            for db in ["PubMed", "Europe PMC", "OpenAlex", "Semantic Scholar", "Crossref", "ERIC", "Scopus", "EMBASE"]:
                if db == "PubMed":
                    n, e = pubmed_count_only(query, date_from, date_to)
                elif db == "Europe PMC":
                    n, e = epmc_count_only(query, date_from, date_to)
                elif db == "OpenAlex":
                    n, e = openalex_count_only(query, date_from, date_to)
                elif db == "Semantic Scholar":
                    n, e = semanticscholar_count_only(query, date_from, date_to)
                elif db == "Crossref":
                    n, e = crossref_count_only(query, date_from, date_to)
                elif db == "ERIC":
                    n, e = eric_count_only(query, date_from, date_to)
                elif db == "Scopus":
                    n, e = scopus_count_only(query, date_from, date_to)
                else:  # EMBASE
                    n, e = embase_count_only(query, date_from, date_to)
                counts[db] = (n, e)

        # Render the count summary prominently — grid of 4 wide
        st.markdown("**Preview counts** (TOTAL matches before any retrieval cap):")
        db_list = list(counts.items())
        rows = [st.columns(4), st.columns(4)]
        for i, (db, (n, e)) in enumerate(db_list):
            col = rows[i // 4][i % 4]
            with col:
                if e:
                    st.metric(db, "—")
                    st.caption(f"⚠️ {e[:60]}")
                else:
                    st.metric(db, f"{n:,}")
                    if n < 30:
                        st.caption("🔴 Very low")
                    elif n < 100:
                        st.caption("🟡 Modest")
                    else:
                        st.caption("🟢 Healthy")

        if all(c[0] < 30 for c in counts.values() if not c[1]):
            st.warning("All databases return very few matches. **Try:** removing the most restrictive AND clause, or use the '🤖 Build query with AI' button at the top to generate a broader query from your topic + criteria.")

    selected_dbs = st.multiselect(
        "Databases to search",
        ["PubMed", "Europe PMC", "OpenAlex", "Semantic Scholar", "Crossref", "ERIC", "Scopus", "EMBASE"],
        default=["PubMed", "Europe PMC", "OpenAlex"],
        help=(
            "**PubMed**: biomedical core, MEDLINE-indexed. "
            "**Europe PMC**: PubMed + Europe-specific + preprints. "
            "**OpenAlex**: 250M+ scholarly works, broad coverage. "
            "**Semantic Scholar**: AI-enhanced semantic search. "
            "**Crossref**: every DOI-registered journal article. "
            "**ERIC**: education research — relevant for accreditation, curriculum. "
            "**Scopus**: requires Elsevier API key (free tier available). "
            "**EMBASE**: requires PAID Elsevier subscription + API key."
        ),
    )

    if search_btn and query.strip():
        progress = st.progress(0, text="Searching...")
        all_results = []
        status_lines = []
        n_dbs = len(selected_dbs)
        step = 100 // max(n_dbs, 1)
        progress_value = 0

        if "PubMed" in selected_dbs:
            progress.progress(min(progress_value + 5, 99), text="Querying PubMed...")
            items, total = search_pubmed(query, date_from, date_to, max_per_db)
            all_results.extend(items)
            status_lines.append(f"PubMed: {len(items)} of {total:,} total")
            progress_value += step

        if "Europe PMC" in selected_dbs:
            progress.progress(min(progress_value + 5, 99), text="Querying Europe PMC...")
            items, total = search_epmc(query, date_from, date_to, max_per_db)
            all_results.extend(items)
            status_lines.append(f"Europe PMC: {len(items)} of {total:,} total")
            progress_value += step

        if "OpenAlex" in selected_dbs:
            progress.progress(min(progress_value + 5, 99), text="Querying OpenAlex...")
            items, total = search_openalex(query, date_from, date_to, max_per_db)
            all_results.extend(items)
            status_lines.append(f"OpenAlex: {len(items)} of {total:,} total")
            progress_value += step

        if "Semantic Scholar" in selected_dbs:
            progress.progress(min(progress_value + 5, 99), text="Querying Semantic Scholar...")
            items, total = search_semanticscholar(query, date_from, date_to, max_per_db)
            all_results.extend(items)
            status_lines.append(f"Semantic Scholar: {len(items)} of {total:,} total")
            progress_value += step

        if "Crossref" in selected_dbs:
            progress.progress(min(progress_value + 5, 99), text="Querying Crossref...")
            items, total = search_crossref(query, date_from, date_to, max_per_db)
            all_results.extend(items)
            status_lines.append(f"Crossref: {len(items)} of {total:,} total")
            progress_value += step

        if "ERIC" in selected_dbs:
            progress.progress(min(progress_value + 5, 99), text="Querying ERIC...")
            items, total = search_eric(query, date_from, date_to, max_per_db)
            all_results.extend(items)
            status_lines.append(f"ERIC: {len(items)} of {total:,} total")
            progress_value += step

        if "Scopus" in selected_dbs:
            progress.progress(min(progress_value + 5, 99), text="Querying Scopus...")
            items, total = search_scopus(query, date_from, date_to, max_per_db)
            all_results.extend(items)
            status_lines.append(f"Scopus: {len(items)} of {total:,} total")
            progress_value += step

        if "EMBASE" in selected_dbs:
            progress.progress(min(progress_value + 5, 99), text="Querying EMBASE...")
            items, total = search_embase(query, date_from, date_to, max_per_db)
            all_results.extend(items)
            status_lines.append(f"EMBASE: {len(items)} of {total:,} total")
            progress_value += step

        deduped = dedup_results(all_results)
        # Initialize stage tracking — all records start at stage 0 (post-dedup, awaiting title screen)
        for i, r in enumerate(deduped):
            r["num"] = i + 1
            r.setdefault("stage", 0)  # 0 = post-dedup, 1 = title-screened, 2 = abstract-screened, 3 = full-text/appraisal-screened
            r.setdefault("title_decision", "Pending")
            r.setdefault("abstract_decision", "Pending")
            r.setdefault("fulltext_decision", "Pending")
            r.setdefault("journal_quality", None)
        st.session_state.results = deduped
        st.session_state.search_status = f"**{len(deduped)} unique results** after deduplication"
        st.session_state.db_status = status_lines
        progress.progress(100, text=f"Done — {len(deduped)} unique results")
        time.sleep(0.5)
        progress.empty()
        st.rerun()

    if st.session_state.search_status:
        st.success(st.session_state.search_status)
        # Render per-database stats with clearer diagnostic styling
        ds_cols = st.columns(len(st.session_state.db_status) if st.session_state.db_status else 1)
        for col, line in zip(ds_cols, st.session_state.db_status):
            with col:
                st.markdown(f"• {line}")
        # If any DB shows "X of Y" where X < Y AND X equals the cap, suggest raising the cap
        capped = []
        for line in st.session_state.db_status:
            # Try to parse "DB: N of M total"
            import re
            m = re.match(r"(\w[\w ]*): (\d[\d,]*) of (\d[\d,]*) total", line)
            if m:
                db_name = m.group(1)
                retrieved = int(m.group(2).replace(",", ""))
                total = int(m.group(3).replace(",", ""))
                if total > retrieved and total > 50:
                    capped.append((db_name, retrieved, total))
        if capped:
            msgs = []
            for db, ret, tot in capped:
                msgs.append(f"**{db}** retrieved {ret:,} of {tot:,} available")
            st.warning(
                "📊 " + " · ".join(msgs)
                + f". Raise **Max per DB** above current value to capture more, or accept the cap if you only want the most recent/relevant. PubMed sorts by relevance by default."
            )

    if st.session_state.results:
        inc = sum(1 for r in st.session_state.results if r["decision"] == "Include")
        exc = sum(1 for r in st.session_state.results if r["decision"] == "Exclude")
        may = sum(1 for r in st.session_state.results if r["decision"] == "Maybe")
        pen = sum(1 for r in st.session_state.results if r["decision"] == "Pending")
        m1, m2, m3, m4 = st.columns(4)
        m1.metric("Total", len(st.session_state.results))
        m2.metric("Include", inc)
        m3.metric("Maybe", may)
        m4.metric("Exclude", exc)
        st.info(f"📋 Switch to the **Screen tab** to apply AI screening against your inclusion criteria.")


# ─── SCREEN TAB ────────────────────────────────────────────────
with tab_screen:
    st.subheader("Staged Screening (PRISMA-ScR)")
    st.caption("Records progress through three stages: **Stage 1 (Title)** → **Stage 2 (Abstract + Journal Quality)** → **Stage 3 (Full-text + Critical Appraisal)**. Only studies passing each stage advance to the next. Final corpus = Stage 3 includes.")

    if not st.session_state.results:
        st.info("Run a search first.")
    else:
        # ─── PRISMA-style stage counters ───
        n_total = len(st.session_state.results)
        n_s1_inc = sum(1 for r in st.session_state.results if r.get("title_decision") == "Include")
        n_s1_exc = sum(1 for r in st.session_state.results if r.get("title_decision") == "Exclude")
        n_s1_pen = sum(1 for r in st.session_state.results if r.get("title_decision", "Pending") == "Pending")
        n_s2_inc = sum(1 for r in st.session_state.results if r.get("abstract_decision") == "Include")
        n_s2_exc = sum(1 for r in st.session_state.results if r.get("abstract_decision") == "Exclude")
        n_s2_pen = sum(1 for r in st.session_state.results if r.get("title_decision") == "Include" and r.get("abstract_decision", "Pending") == "Pending")
        n_s3_inc = sum(1 for r in st.session_state.results if r.get("fulltext_decision") == "Include")
        n_s3_exc = sum(1 for r in st.session_state.results if r.get("fulltext_decision") == "Exclude")
        n_s3_pen = sum(1 for r in st.session_state.results if r.get("abstract_decision") == "Include" and r.get("fulltext_decision", "Pending") == "Pending")

        progress_cols = st.columns(4)
        progress_cols[0].metric("After dedup", n_total)
        progress_cols[1].metric("Stage 1 included", n_s1_inc, delta=f"-{n_s1_exc} excluded")
        progress_cols[2].metric("Stage 2 included", n_s2_inc, delta=f"-{n_s2_exc} excluded")
        progress_cols[3].metric("Stage 3 included", n_s3_inc, delta=f"-{n_s3_exc} excluded")

        st.divider()

        # ─── Choose stage to work in ───
        stage_choice = st.radio(
            "Active stage",
            ["Stage 1 – Title screening", "Stage 2 – Abstract screening", "Stage 3 – Full-text + Critical Appraisal"],
            horizontal=True,
            key="active_stage",
        )

        if "Stage 1" in stage_choice:
            stage_num = 1
            field = "title_decision"
            queue = [r for r in st.session_state.results if r.get("title_decision", "Pending") == "Pending"]
            other_field_filter = lambda r: True  # all records eligible
        elif "Stage 2" in stage_choice:
            stage_num = 2
            field = "abstract_decision"
            queue = [r for r in st.session_state.results
                     if r.get("title_decision") == "Include" and r.get("abstract_decision", "Pending") == "Pending"]
            other_field_filter = lambda r: r.get("title_decision") == "Include"
        else:
            stage_num = 3
            field = "fulltext_decision"
            queue = [r for r in st.session_state.results
                     if r.get("abstract_decision") == "Include" and r.get("fulltext_decision", "Pending") == "Pending"]
            other_field_filter = lambda r: r.get("abstract_decision") == "Include"

        st.caption(f"**{len(queue)} records pending in this stage.**")

        # ─── Bulk AI actions per stage ───
        bulk_cols = st.columns([2, 2, 2])
        with bulk_cols[0]:
            if stage_num == 1:
                btn_label = f"🤖 AI title-screen all pending ({len(queue)})"
            elif stage_num == 2:
                btn_label = f"🤖 AI abstract-screen all pending ({len(queue)})"
            else:
                btn_label = f"🤖 Appraise all pending ({len(queue)})"

            if st.button(btn_label, type="primary", disabled=not client or len(queue) == 0):
                if client:
                    progress = st.progress(0, text="Starting...")
                    # Stage 1 & 2 use batch calls (up to 20/10 articles per API call)
                    if stage_num == 1:
                        BATCH = 20
                        processed = 0
                        for batch_start in range(0, len(queue), BATCH):
                            batch = queue[batch_start: batch_start + BATCH]
                            progress.progress(
                                min((batch_start + BATCH) / len(queue), 1.0),
                                text=f"Title-screening {batch_start+1}–{min(batch_start+BATCH, len(queue))} of {len(queue)}…",
                            )
                            batch_results, err = batch_screen_title(batch)
                            if batch_results:
                                for art, result in zip(batch, batch_results):
                                    for rr in st.session_state.results:
                                        if rr["id"] == art["id"]:
                                            rr[field] = result.get("decision", "Maybe")
                                            rr["rationale"] = result.get("rationale", "")
                                            rr["confidence"] = result.get("confidence", "")
                                            rr["decision"] = result.get("decision", "Maybe")
                                            rr["stage"] = stage_num
                                            break
                            else:
                                # Fall back to single-article calls for this batch
                                for art in batch:
                                    result, serr = screen_article_title_only(art)
                                    for rr in st.session_state.results:
                                        if rr["id"] == art["id"]:
                                            rr[field] = result.get("decision", "Maybe") if result else "Maybe"
                                            rr["rationale"] = result.get("rationale", "") if result else f"AI error: {serr}"
                                            rr["confidence"] = result.get("confidence", "") if result else ""
                                            rr["decision"] = rr[field]
                                            rr["stage"] = stage_num
                                            break
                            time.sleep(0.3)

                    elif stage_num == 2:
                        BATCH = 10
                        for batch_start in range(0, len(queue), BATCH):
                            batch = queue[batch_start: batch_start + BATCH]
                            progress.progress(
                                min((batch_start + BATCH) / len(queue), 1.0),
                                text=f"Abstract-screening {batch_start+1}–{min(batch_start+BATCH, len(queue))} of {len(queue)}…",
                            )
                            batch_results, err = batch_screen_abstract(batch)
                            if batch_results:
                                for art, result in zip(batch, batch_results):
                                    jq = check_journal_quality(art.get("journal", ""))
                                    if jq["flag"] == "critical":
                                        result["decision"] = "Exclude"
                                        result["rationale"] = (result.get("rationale", "") + f" | Journal quality: {jq['reason']}").strip(" |")
                                    for rr in st.session_state.results:
                                        if rr["id"] == art["id"]:
                                            rr["journal_quality"] = jq
                                            rr[field] = result.get("decision", "Maybe")
                                            rr["rationale"] = result.get("rationale", "")
                                            rr["confidence"] = result.get("confidence", "")
                                            rr["decision"] = result.get("decision", "Maybe")
                                            rr["tier"] = result.get("tier", rr.get("tier", ""))
                                            rr["kp"] = result.get("kp", rr.get("kp", "N/A"))
                                            rr["stage"] = stage_num
                                            break
                            else:
                                # Fall back to single-article calls for this batch
                                for art in batch:
                                    result, serr = screen_article(art)
                                    jq = check_journal_quality(art.get("journal", ""))
                                    if result and jq["flag"] == "critical":
                                        result["decision"] = "Exclude"
                                        result["rationale"] = (result.get("rationale", "") + f" | Journal quality: {jq['reason']}").strip(" |")
                                    for rr in st.session_state.results:
                                        if rr["id"] == art["id"]:
                                            rr["journal_quality"] = jq
                                            rr[field] = result.get("decision", "Maybe") if result else "Maybe"
                                            rr["rationale"] = result.get("rationale", "") if result else f"AI error: {serr}"
                                            rr["confidence"] = result.get("confidence", "") if result else ""
                                            rr["decision"] = rr[field]
                                            if result:
                                                rr["tier"] = result.get("tier", rr.get("tier", ""))
                                                rr["kp"] = result.get("kp", rr.get("kp", "N/A"))
                                            rr["stage"] = stage_num
                                            break
                            time.sleep(0.3)

                    else:
                        # Stage 3: Critical appraisal — one at a time (complex per-study output)
                        for i, art in enumerate(queue):
                            progress.progress((i + 1) / len(queue), text=f"{i+1}/{len(queue)}: {art['title'][:50]}...")
                            result, err = appraise_study(art)
                            if result:
                                st.session_state.appraisals[art["id"]] = result
                                decision = result.get("recommendation", "Maybe")
                                rationale = f"{result.get('overall_rating','')} quality — {result.get('rationale','')}"
                                result = {"decision": decision, "rationale": rationale, "confidence": "high"}
                            for rr in st.session_state.results:
                                if rr["id"] == art["id"]:
                                    rr[field] = result.get("decision", "Maybe") if result else "Maybe"
                                    rr["rationale"] = result.get("rationale", "") if result else f"AI error: {err}"
                                    rr["confidence"] = result.get("confidence", "") if result else ""
                                    rr["decision"] = rr[field]
                                    rr["stage"] = stage_num
                                    break
                            time.sleep(0.3)

                    progress.empty()
                    st.success(f"Stage {stage_num} screening complete.")
                    st.rerun()
        with bulk_cols[1]:
            if stage_num == 3 and st.button("➕ Add Stage 3 includes to corpus", type="secondary"):
                existing = {c["id"] for c in st.session_state.corpus}
                added = 0
                for r in st.session_state.results:
                    if r.get("fulltext_decision") == "Include" and r["id"] not in existing:
                        st.session_state.corpus.append(dict(r))
                        added += 1
                st.success(f"Added {added} studies to the final corpus.")
                st.rerun()
        with bulk_cols[2]:
            filter_db = st.selectbox(
                "Database filter",
                ["All", "PubMed", "Europe PMC", "OpenAlex", "Semantic Scholar", "Crossref", "ERIC", "Scopus", "EMBASE"],
                key="filter_db_stage",
            )

        st.divider()

        # ─── Records list — only show records eligible for current stage ───
        records_in_stage = [r for r in st.session_state.results if other_field_filter(r)]
        if filter_db != "All":
            records_in_stage = [r for r in records_in_stage if r["db"] == filter_db]

        # Decision filter for current stage
        filter_dec = st.selectbox(
            f"Stage {stage_num} decision filter",
            ["All", "Pending", "Include", "Maybe", "Exclude"],
            key=f"filter_dec_s{stage_num}",
        )
        if filter_dec != "All":
            records_in_stage = [r for r in records_in_stage if r.get(field, "Pending") == filter_dec]

        st.caption(f"Showing {len(records_in_stage)} records at Stage {stage_num}")

        for r in records_in_stage:
            current_decision = r.get(field, "Pending")
            with st.container(border=True):
                col_main, col_btns = st.columns([5, 2])
                with col_main:
                    st.markdown(f"**#{r['num']}** · [{r['title']}]({r.get('url','')})")
                    st.caption(f"{r['authors']} · *{r['journal']}* {r['year']} · {r['db']}")
                    # Journal quality indicator (Stage 2 and later)
                    if stage_num >= 2 and r.get("journal_quality"):
                        jq = r["journal_quality"]
                        emoji = {"good": "🟢", "warning": "🟡", "critical": "🔴"}.get(jq["flag"], "")
                        st.caption(f"{emoji} **Journal quality:** {jq['reason']}")
                with col_btns:
                    dec_cols = st.columns(3)
                    for j, d in enumerate(["Include", "Maybe", "Exclude"]):
                        with dec_cols[j]:
                            current = current_decision == d
                            if st.button(d, key=f"dec_s{stage_num}_{r['id']}_{d}", type="primary" if current else "secondary", use_container_width=True):
                                for rr in st.session_state.results:
                                    if rr["id"] == r["id"]:
                                        rr[field] = d
                                        rr["decision"] = d
                                        rr["stage"] = stage_num
                                        break
                                st.rerun()

                # Tier & Kirkpatrick (only at Stage 2+, when we know more)
                if stage_num >= 2:
                    tk_cols = st.columns([3, 3, 2, 2])
                    with tk_cols[0]:
                        tier_opts = [""] + st.session_state.config["tiers"]
                        current_tier_idx = tier_opts.index(r.get("tier", "")) if r.get("tier", "") in tier_opts else 0
                        new_tier = st.selectbox("Tier", tier_opts, index=current_tier_idx, key=f"tier_s{stage_num}_{r['id']}", label_visibility="collapsed")
                        if new_tier != r.get("tier", ""):
                            for rr in st.session_state.results:
                                if rr["id"] == r["id"]:
                                    rr["tier"] = new_tier
                                    break
                    with tk_cols[1]:
                        kp_opts = st.session_state.config["kp_levels"]
                        current_kp_idx = kp_opts.index(r.get("kp", "N/A")) if r.get("kp", "N/A") in kp_opts else 0
                        new_kp = st.selectbox("KP", kp_opts, index=current_kp_idx, key=f"kp_s{stage_num}_{r['id']}", label_visibility="collapsed")
                        if new_kp != r.get("kp", "N/A"):
                            for rr in st.session_state.results:
                                if rr["id"] == r["id"]:
                                    rr["kp"] = new_kp
                                    break
                    with tk_cols[2]:
                        if st.button("🤖 AI", key=f"aiscr_s{stage_num}_{r['id']}", disabled=not client, use_container_width=True):
                            with st.spinner("Working..."):
                                if stage_num == 2:
                                    result, err = screen_article(r)
                                    jq = check_journal_quality(r.get("journal", ""))
                                    for rr in st.session_state.results:
                                        if rr["id"] == r["id"]:
                                            rr["journal_quality"] = jq
                                            break
                                elif stage_num == 3:
                                    ap, err = appraise_study(r)
                                    if ap:
                                        st.session_state.appraisals[r["id"]] = ap
                                        result = {"decision": ap.get("recommendation", "Maybe"),
                                                  "rationale": f"{ap.get('overall_rating','')} quality — {ap.get('rationale','')}",
                                                  "confidence": "high"}
                                    else:
                                        result = None
                                if result:
                                    for rr in st.session_state.results:
                                        if rr["id"] == r["id"]:
                                            rr[field] = result.get("decision", "Maybe")
                                            rr["decision"] = result.get("decision", "Maybe")
                                            rr["rationale"] = result.get("rationale", "")
                                            rr["confidence"] = result.get("confidence", "")
                                            rr["stage"] = stage_num
                                            break
                                    st.rerun()
                                else:
                                    st.error(f"Error: {err}")
                    with tk_cols[3]:
                        if stage_num == 3 and current_decision == "Include":
                            if st.button("➕ Corpus", key=f"corp_s{stage_num}_{r['id']}", use_container_width=True):
                                if not any(c["id"] == r["id"] for c in st.session_state.corpus):
                                    st.session_state.corpus.append(dict(r))
                                st.rerun()

                # Rationale
                if r.get("rationale"):
                    conf_emoji = {"high": "🟢", "medium": "🟡", "low": "🔴"}.get(r.get("confidence", ""), "")
                    st.caption(f"*{conf_emoji} {r['rationale']}*")

                # Show appraisal details if available (Stage 3)
                if stage_num == 3 and r["id"] in st.session_state.appraisals:
                    ap = st.session_state.appraisals[r["id"]]
                    with st.expander(f"📋 Appraisal: {ap.get('tool','')} → {ap.get('overall_rating','')}"):
                        rating_color = {"High": "🟢", "Moderate": "🟡", "Low": "🔴"}.get(ap.get("overall_rating", ""), "")
                        st.markdown(f"**Overall:** {rating_color} {ap.get('overall_rating','')} — {ap.get('rationale','')}")
                        st.markdown(f"**Recommendation:** {ap.get('recommendation','')}")
                        if ap.get("exclusion_reason"):
                            st.warning(f"Exclusion reason: {ap['exclusion_reason']}")
                        st.markdown("**Item scores:**")
                        for s in ap.get("item_scores", []):
                            r_emoji = {"Yes": "✅", "No": "❌", "Unclear": "❓", "NA": "—"}.get(s.get("rating", ""), "")
                            st.markdown(f"- {r_emoji} **{s.get('rating','')}** — {s.get('item','')}: {s.get('note','')}")


# ─── CORPUS TAB ────────────────────────────────────────────────
with tab_corpus:
    st.subheader(f"Corpus — {len(st.session_state.corpus)} studies")

    if not st.session_state.corpus:
        st.info("Mark articles as Include in the Screen tab, then add them to the corpus.")
    else:
        c1, c2 = st.columns(2)
        with c1:
            n_unextracted = sum(1 for c in st.session_state.corpus if not c.get("extraction"))
            if st.button(f"🤖 AI extract all ({n_unextracted} remaining)", type="primary", disabled=not client or n_unextracted == 0):
                if client:
                    progress = st.progress(0, text="Extracting...")
                    to_extract = [c for c in st.session_state.corpus if not c.get("extraction")]
                    for i, art in enumerate(to_extract):
                        progress.progress((i + 1) / len(to_extract), text=f"Extracting {i+1}/{len(to_extract)}")
                        result, err = extract_article(art)
                        for c in st.session_state.corpus:
                            if c["id"] == art["id"]:
                                c["extraction"] = result if result else {"error": err}
                                break
                        time.sleep(0.4)
                    progress.empty()
                    st.success(f"Extracted {len(to_extract)} studies.")
                    st.rerun()
        with c2:
            # CSV export
            rows = []
            for c in st.session_state.corpus:
                e = c.get("extraction") or {}
                rows.append({
                    "PMID/DOI": c["pmid"],
                    "Title": c["title"],
                    "Authors": c["authors"],
                    "Journal": c["journal"],
                    "Year": c["year"],
                    "Database": c["db"],
                    "Decision": c["decision"],
                    "Confidence": c.get("confidence", ""),
                    "Tier": c.get("tier", ""),
                    "Kirkpatrick": c.get("kp", ""),
                    "Screening rationale": c.get("rationale", ""),
                    "Country": e.get("country", ""),
                    "Setting": e.get("setting", ""),
                    "Design": e.get("design", ""),
                    "Sample": e.get("sample", ""),
                    "AI tool": e.get("aiTool", ""),
                    "Educational outcome": e.get("educationalOutcome", ""),
                    "Key finding": e.get("keyFinding", ""),
                    "Main limitation": e.get("mainLimitation", ""),
                    "Relevance": e.get("relevance", ""),
                    "URL": c["url"],
                })
            df = pd.DataFrame(rows)
            st.download_button(
                "📥 Export corpus CSV",
                df.to_csv(index=False),
                file_name=f"corpus_{datetime.now().strftime('%Y%m%d_%H%M')}.csv",
                mime="text/csv",
                type="primary",
                use_container_width=True,
            )

        for c in st.session_state.corpus:
            with st.container(border=True):
                col_main, col_btn = st.columns([6, 1])
                with col_main:
                    st.markdown(f"**[{c['title']}]({c['url']})**")
                    badges = c["db"]
                    if c.get("tier"):
                        badges += f" · {c['tier']}"
                    if c.get("kp") and c["kp"] != "N/A":
                        badges += f" · {c['kp']}"
                    st.caption(f"{c['authors']} · *{c['journal']}* {c['year']} · {badges}")
                    if c.get("rationale"):
                        st.caption(f"*{c['rationale']}*")
                with col_btn:
                    if st.button("Remove", key=f"rm_{c['id']}", use_container_width=True):
                        st.session_state.corpus = [x for x in st.session_state.corpus if x["id"] != c["id"]]
                        st.rerun()

                if c.get("extraction"):
                    e = c["extraction"]
                    if "error" in e:
                        st.error(f"Extraction error: {e['error']}")
                    else:
                        ec1, ec2 = st.columns(2)
                        with ec1:
                            for k in ["country", "setting", "design", "sample"]:
                                if e.get(k):
                                    st.markdown(f"**{k.capitalize()}:** {e[k]}")
                        with ec2:
                            for k in ["aiTool", "educationalOutcome", "relevance"]:
                                if e.get(k):
                                    label = {"aiTool": "AI tool", "educationalOutcome": "Outcome", "relevance": "Relevance"}[k]
                                    st.markdown(f"**{label}:** {e[k]}")
                        if e.get("keyFinding"):
                            st.markdown(f"**Key finding:** {e['keyFinding']}")
                        if e.get("mainLimitation"):
                            st.markdown(f"**Limitation:** {e['mainLimitation']}")


# ─── SYNTHESIS TAB ────────────────────────────────────────────
with tab_synth:
    st.subheader("Structured corpus synthesis")
    n_corpus = len(st.session_state.corpus)
    has_style = bool(st.session_state.config.get("writing_style_sample", "").strip())
    n_extracted = sum(1 for c in st.session_state.corpus if c.get("extraction"))
    has_themes = bool(st.session_state.themes)

    cap_bits = [f"Corpus n={n_corpus}"]
    cap_bits.append("✓ writing style loaded" if has_style else "⚠ no writing style set")
    cap_bits.append(f"{n_extracted}/{n_corpus} extracted")
    cap_bits.append(f"✓ {len(st.session_state.themes)} themes discovered" if has_themes else "tier-based")
    st.caption(" · ".join(cap_bits))

    if n_corpus == 0:
        st.info("Add studies to the corpus first, then optionally run AI extraction for richer geographic/design tables.")
    else:
        # ─── Theme discovery panel ───
        with st.expander("🧩 Discover emergent themes (inductive thematic analysis)", expanded=has_themes):
            st.caption("Run AI thematic analysis on extracted study findings. Themes emerge from the data inductively, replacing pre-defined tiers in the synthesis structure.")

            tc1, tc2 = st.columns([2, 1])
            with tc1:
                if st.button(
                    "🧩 Discover themes from extractions",
                    type="secondary",
                    disabled=not client or n_extracted < 3,
                    help=f"{n_extracted} extracted studies available (need ≥3)",
                ):
                    if client:
                        with st.spinner(f"Analyzing {n_extracted} extracted studies..."):
                            themes, err = discover_themes(st.session_state.corpus)
                            if themes:
                                st.session_state.themes = themes
                                st.session_state.synth_data = None  # invalidate prior synthesis
                                st.rerun()
                            else:
                                st.error(f"Error: {err}")
            with tc2:
                if has_themes and st.button("Clear themes", type="secondary"):
                    st.session_state.themes = None
                    st.session_state.synth_data = None
                    st.rerun()

            if has_themes:
                st.markdown("**Discovered themes:**")
                for i, t in enumerate(st.session_state.themes, 1):
                    with st.container(border=True):
                        st.markdown(f"**Theme {i} — {t.get('name','')}**")
                        st.markdown(t.get("definition", ""))
                        if t.get("supporting_studies"):
                            st.caption(f"Supporting studies: {', '.join(t['supporting_studies'][:8])}")
                        if t.get("theoretical_link"):
                            st.caption(f"🎯 Theoretical link: {t['theoretical_link']}")
                        if t.get("tensions"):
                            st.caption(f"⚠️ Tensions: {t['tensions']}")

        st.divider()

        c1, c2 = st.columns([2, 1])
        with c1:
            synth_label = "✨ Generate synthesis (theme-based)" if has_themes else "✨ Generate synthesis (tier-based)"
            if st.button(synth_label, type="primary", disabled=not client):
                if client:
                    with st.spinner("Computing statistics and generating narrative (this may take 60–90s)..."):
                        result, err = synthesise_corpus(
                            st.session_state.corpus,
                            st.session_state.config,
                            themes=st.session_state.themes,
                        )
                        if result:
                            st.session_state.synth_data = result
                            st.session_state.synth_text = result["narrative"]
                            st.session_state.synth_chat = []  # reset chat on new synthesis
                        else:
                            st.error(f"Error: {err}")
        with c2:
            if st.session_state.synth_data:
                # Build downloadable markdown
                stats = st.session_state.synth_data["stats"]
                narrative = st.session_state.synth_data["narrative"]
                sections = [s.strip() for s in narrative.split("===SECTION===")]
                use_themes_here = st.session_state.synth_data.get("use_themes", False)
                if use_themes_here:
                    section_titles = [
                        "1. Corpus Overview",
                        "2. Thematic Landscape",
                        "3. Outcome Evidence Profile",
                        "4. Geographic and Contextual Variation",
                        "5. Cross-Theme Patterns",
                        "6. Comparison with Available Literature",
                        "7. Integration Across Themes",
                        "8. Argument and Implications",
                        "9. Conclusion",
                    ]
                else:
                    section_titles = [
                        "1. Corpus Overview",
                        "2. Tier Distribution and Evidence Weight",
                        "3. Kirkpatrick Outcome Ceiling",
                        "4. Geographic and Contextual Variation",
                        "5. Cross-Tier Patterns",
                        "6. Comparison with Available Literature",
                        "7. Integration Across Tiers",
                        "8. Argument and Implications",
                        "9. Conclusion",
                    ]

                # Compute PRISMA numbers for the markdown export
                prisma = compute_prisma_numbers(st.session_state.results, st.session_state.corpus)
                prisma_md = ""
                if prisma:
                    prisma_md = (
                        "## PRISMA-ScR Flow\n\n"
                        f"- Records identified across databases: "
                        + ", ".join([f"{db} n={n:,}" for db, n in sorted(prisma['by_db'].items(), key=lambda x: -x[1])])
                        + f"\n- Records after duplicates removed: n = {prisma['n_after_dedup']:,}\n"
                        f"- Records screened: n = {prisma['n_screened']:,}\n"
                        f"- Records excluded at screening: n = {prisma['n_excluded_screening']:,}\n"
                    )
                    if prisma['exclusion_reasons']:
                        prisma_md += "  - Reasons:\n"
                        for reason, count in sorted(prisma['exclusion_reasons'].items(), key=lambda x: -x[1]):
                            prisma_md += f"    - {reason}: n = {count}\n"
                    prisma_md += (
                        f"- Records uncertain (Maybe / deferred): n = {prisma['n_maybe']:,}\n"
                        f"- Studies included in synthesis: n = {prisma['n_corpus']:,}\n\n"
                    )

                md_parts = [f"# Synthesis: {st.session_state.config['research_topic']}\n"]
                md_parts.append(f"_Generated {datetime.now().strftime('%Y-%m-%d %H:%M')} · n={stats['n']} studies_\n\n")
                if prisma_md:
                    md_parts.append(prisma_md)

                # Overview table (Table 1)
                overview_md = build_overview_table_md(st.session_state.corpus, st.session_state.config)
                if overview_md:
                    md_parts.append("## Table 1. Overview of Included Studies\n")
                    md_parts.append(overview_md + "\n\n")

                for i, title in enumerate(section_titles):
                    md_parts.append(f"## {title}\n")
                    # Add table for relevant sections
                    if i == 0:
                        md_parts.append(f"**Total studies:** {stats['n']}  ·  **Year range:** {stats['year_range']}\n\n")
                        md_parts.append("### Database sourcing\n" + build_db_table(stats) + "\n\n")
                    elif i == 1:
                        md_parts.append("### Tier distribution\n" + build_tier_table(stats, st.session_state.config) + "\n\n")
                    elif i == 2:
                        md_parts.append("### Kirkpatrick level distribution\n" + build_kp_table(stats) + "\n\n")
                    elif i == 3:
                        geo_t = build_geo_table(stats)
                        if geo_t:
                            md_parts.append("### Geographic distribution\n" + geo_t + "\n\n")
                    if i < len(sections):
                        md_parts.append(sections[i] + "\n\n")

                # ─── Manchester-Harvard References ───
                refs = build_reference_list(st.session_state.corpus)
                if refs:
                    md_parts.append("## References\n")
                    md_parts.append("_Manchester-Harvard style_\n\n")
                    for ref in refs:
                        md_parts.append(f"{ref}\n\n")

                full_md = "".join(md_parts)
                st.download_button(
                    "📥 Download (.md)",
                    full_md,
                    file_name=f"synthesis_{datetime.now().strftime('%Y%m%d_%H%M')}.md",
                    mime="text/markdown",
                    use_container_width=True,
                )

                # ─── Word document download ───
                try:
                    word_bio = build_word_document(
                        st.session_state.corpus,
                        stats,
                        narrative,
                        prisma,
                        st.session_state.config,
                    )
                    st.download_button(
                        "📄 Download (.docx)",
                        data=word_bio,
                        file_name=f"synthesis_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                    )
                except ImportError:
                    st.caption("⚠️ Install `python-docx` to enable Word export.")
                except Exception as e:
                    st.caption(f"⚠️ Word export failed: {e}")

        # ─── Render structured synthesis output ────────────────────
        if st.session_state.synth_data:
            stats = st.session_state.synth_data["stats"]
            narrative = st.session_state.synth_data["narrative"]
            sections = [s.strip() for s in narrative.split("===SECTION===")]

            # Pad sections list so missing ones don't crash
            while len(sections) < 9:
                sections.append("(Section narrative not generated — try regenerating.)")

            st.divider()

            # ───────────── PRISMA-ScR FLOWCHART (at the top) ─────────────
            st.markdown("### PRISMA-ScR Flow Diagram")
            st.caption("Records identified, screened, excluded with reasons, and included in the synthesis.")
            prisma = compute_prisma_numbers(st.session_state.results, st.session_state.corpus)
            if prisma:
                prisma_code = build_prisma_flowchart(prisma)
                render_mermaid(prisma_code, height=720)
                with st.expander("PRISMA flow diagram source (Mermaid)", expanded=False):
                    st.code(prisma_code, language="mermaid")
                with st.expander("PRISMA numbers in tabular form", expanded=False):
                    rows = ["| Stage | n |", "|---|---|"]
                    for db, n in sorted(prisma["by_db"].items(), key=lambda x: -x[1]):
                        rows.append(f"| Records identified — {db} | {n:,} |")
                    rows.append(f"| Records after duplicates removed | {prisma['n_after_dedup']:,} |")
                    rows.append(f"| Records screened | {prisma['n_screened']:,} |")
                    rows.append(f"| Records excluded at screening | {prisma['n_excluded_screening']:,} |")
                    for reason, count in sorted(prisma["exclusion_reasons"].items(), key=lambda x: -x[1]):
                        rows.append(f"|  &nbsp;&nbsp;– {reason} | {count} |")
                    rows.append(f"| Records uncertain (Maybe) | {prisma['n_maybe']:,} |")
                    rows.append(f"| Studies included | {prisma['n_corpus']:,} |")
                    st.markdown("\n".join(rows))
            else:
                st.info("Run a search and screen some results to populate the PRISMA flowchart.")

            st.divider()

            # ──────────────────── TABLE 1: OVERVIEW OF INCLUDED STUDIES ────────────────────
            st.markdown("### Table 1. Overview of Included Studies")
            st.caption(f"All {stats['n']} included studies grouped by tier, ordered by year ascending. Empty cells indicate the field was not extracted or not reported.")
            overview_df = build_overview_table_df(st.session_state.corpus, st.session_state.config)
            if overview_df is not None and not overview_df.empty:
                st.dataframe(
                    overview_df,
                    hide_index=True,
                    use_container_width=True,
                    column_config={
                        "#": st.column_config.NumberColumn(width="small"),
                        "Tier": st.column_config.TextColumn(width="medium"),
                        "Study": st.column_config.TextColumn(width="small"),
                        "Country/Setting": st.column_config.TextColumn(width="small"),
                        "Design (Sample)": st.column_config.TextColumn(width="medium"),
                        "AI Tool": st.column_config.TextColumn(width="small"),
                        "Educational Outcome": st.column_config.TextColumn(width="medium"),
                        "Key Finding": st.column_config.TextColumn(width="large"),
                        "Limitation": st.column_config.TextColumn(width="medium"),
                        "Kirkpatrick": st.column_config.TextColumn(width="small"),
                        "PMID/DOI": st.column_config.TextColumn(width="small"),
                    },
                )
                n_extracted_now = sum(1 for c in st.session_state.corpus if c.get("extraction"))
                if n_extracted_now < stats["n"]:
                    st.caption(f"⚠️ {stats['n'] - n_extracted_now} of {stats['n']} studies have not been AI-extracted. Run AI extraction in the Corpus tab to populate Country, Design, AI Tool, and Key Finding fields.")
            else:
                st.info("No corpus available to display.")

            st.divider()

            # ──────────────────── RESULTS SECTIONS ────────────────────
            st.markdown("## Results")

            # Section 1: Corpus Overview
            st.markdown("### 1. Corpus Overview")
            ov_cols = st.columns(4)
            ov_cols[0].metric("Total studies", stats["n"])
            ov_cols[1].metric("Year range", stats["year_range"])
            ov_cols[2].metric("Databases", len(stats["db_counts"]))
            ov_cols[3].metric("Tiers used", sum(1 for c in stats["tier_counts"].values() if c > 0))
            st.markdown("**Database sourcing**")
            st.markdown(build_db_table(stats))
            st.markdown(sections[0])

            st.divider()

            # Section 2: Tier Distribution OR Thematic Landscape
            use_themes_display = st.session_state.synth_data.get("use_themes", False)
            if use_themes_display:
                st.markdown("### 2. Thematic Landscape")
                themes_local = st.session_state.synth_data.get("themes") or []
                # Build a theme summary table
                if themes_local:
                    rows = ["| # | Theme | Supporting studies |", "|---|---|---|"]
                    for i, t in enumerate(themes_local, 1):
                        n_studies = len(t.get("supporting_studies", []))
                        rows.append(f"| {i} | {t.get('name','')} | {n_studies} |")
                    st.markdown("\n".join(rows))
            else:
                st.markdown("### 2. Tier Distribution and Evidence Weight")
                st.markdown(build_tier_table(stats, st.session_state.config))
            st.markdown(sections[1])

            st.divider()

            # Section 3: Kirkpatrick Ceiling OR Outcome Evidence Profile
            if use_themes_display:
                st.markdown("### 3. Outcome Evidence Profile")
                st.markdown(build_kp_table(stats))
            else:
                st.markdown("### 3. Kirkpatrick Outcome Ceiling")
                st.markdown(build_kp_table(stats))
            st.markdown(sections[2])

            st.divider()

            # Section 4: Geographic Spread
            st.markdown("### 4. Geographic and Contextual Variation")
            geo_t = build_geo_table(stats)
            if geo_t:
                st.markdown(geo_t)
            else:
                st.info("No geographic data available — run AI extraction on corpus studies for richer geographic detail.")
            st.markdown(sections[3])

            st.divider()

            # Section 5: Cross-Tier or Cross-Theme Patterns
            if use_themes_display:
                st.markdown("### 5. Cross-Theme Patterns")
            else:
                st.markdown("### 5. Cross-Tier Patterns")
            st.markdown(sections[4])

            st.divider()

            # ──────────────────── DISCUSSION SECTIONS ────────────────────
            st.markdown("## Discussion")

            # Section 6: Comparison with Available Literature
            st.markdown("### 6. Comparison with Available Literature")
            st.markdown(sections[5])

            st.divider()

            # Section 7: Integration Across Tiers / Themes
            if use_themes_display:
                st.markdown("### 7. Integration Across Themes")
            else:
                st.markdown("### 7. Integration Across Tiers")
            st.markdown(sections[6])

            st.divider()

            # Section 8: Argument and Implications
            st.markdown("### 8. Argument and Implications")
            st.markdown(sections[7])

            st.divider()

            # Section 9: Conclusion
            st.markdown("## Conclusion")
            st.markdown(sections[8])

            st.divider()

            # ──────────────────── CONVERSATIONAL REFINEMENT ────────────────────
            st.markdown("## 💬 Refine the synthesis")
            st.caption("Ask follow-up questions or request rewrites. The chat sees the current synthesis + your corpus. Suggested prompts: *Rewrite Section 8 with a stronger argument*, *Add a counter-position to the conclusion*, *Make the comparison section less hedged*, *Connect Theme 1 and Theme 3 more explicitly*.")

            # Display chat history
            for turn in st.session_state.synth_chat:
                with st.chat_message("user"):
                    st.markdown(turn["user"])
                with st.chat_message("assistant"):
                    st.markdown(turn["assistant"])

            # Chat input
            user_msg = st.chat_input("Ask a question or request a rewrite…", disabled=not client)
            if user_msg and client:
                with st.chat_message("user"):
                    st.markdown(user_msg)
                with st.chat_message("assistant"):
                    with st.spinner("Refining..."):
                        response, err = refine_synthesis(
                            current_narrative=st.session_state.synth_data["narrative"],
                            chat_history=st.session_state.synth_chat,
                            user_message=user_msg,
                            corpus=st.session_state.corpus,
                        )
                        if response:
                            st.markdown(response)
                            st.session_state.synth_chat.append({"user": user_msg, "assistant": response})
                        else:
                            st.error(f"Error: {err}")

            if st.session_state.synth_chat:
                if st.button("🗑️ Clear chat history"):
                    st.session_state.synth_chat = []
                    st.rerun()


# ─── TRENDS TAB ─────────────────────────────────────────────────
with tab_trends:
    st.subheader("📈 Recent trends in major journals")
    st.caption("Pull recent papers from a shortlist of journals (via PubMed) and use AI to surface emerging trends. Useful for scoping out a research area or finding angles a manuscript could position itself against.")

    # Journal selection
    selected_journals = st.multiselect(
        "Journals to include",
        DENTAL_JOURNAL_SHORTLIST,
        default=DENTAL_JOURNAL_SHORTLIST[:5],
        help="Default shortlist is dental/medical-education-oriented. Add custom journals below if needed.",
    )

    custom_journals_text = st.text_area(
        "Additional journals (one per line, exact PubMed journal title)",
        value="",
        height=80,
        help="E.g., 'Frontiers in Education' or 'Computers & Education'. Use the exact title as it appears in PubMed.",
    )
    if custom_journals_text.strip():
        custom_journals = [j.strip() for j in custom_journals_text.split("\n") if j.strip()]
        selected_journals = selected_journals + custom_journals

    tc1, tc2, tc3 = st.columns([1, 1, 1])
    with tc1:
        years_back = st.number_input("Years back", min_value=1, max_value=10, value=2)
    with tc2:
        max_per_journal = st.number_input("Max papers per journal", min_value=10, max_value=200, value=30, step=10)
    with tc3:
        st.write("")
        st.write("")
        trends_btn = st.button("🔍 Pull recent papers", type="primary", use_container_width=True, disabled=not selected_journals)

    if trends_btn:
        with st.spinner(f"Fetching from {len(selected_journals)} journals..."):
            papers, per_journal = search_journals_recent(selected_journals, years_back, max_per_journal)
            st.session_state.trends_papers = papers
            # Show per-journal counts
            count_lines = []
            for j in selected_journals:
                d = per_journal.get(j, {"retrieved": 0, "total": 0})
                count_lines.append(f"  • {j}: {d['retrieved']:,} retrieved (of {d['total']:,} total)")
            st.success(f"Pulled {len(papers)} unique papers across {len(selected_journals)} journals.")
            st.code("\n".join(count_lines), language="text")

    if st.session_state.trends_papers:
        st.markdown(f"### Recent corpus: {len(st.session_state.trends_papers)} papers")

        # Topic filter input
        topic_filter = st.text_input(
            "Optional: filter papers by keyword (case-insensitive, searches title + abstract)",
            value="",
            placeholder="e.g., artificial intelligence, accreditation, generative",
        )
        filtered = st.session_state.trends_papers
        if topic_filter.strip():
            kw = topic_filter.lower().strip()
            filtered = [
                p for p in st.session_state.trends_papers
                if kw in (p.get("title", "") + " " + p.get("abstract", "")).lower()
            ]
            st.caption(f"Filter active: {len(filtered)} of {len(st.session_state.trends_papers)} papers match")

        # Preview table
        if filtered:
            df = pd.DataFrame([{
                "Year": p.get("year", ""),
                "Authors": (p.get("authors", "") or "")[:60],
                "Title": (p.get("title", "") or "")[:120],
                "Journal": p.get("journal", ""),
                "Abstract": "✓" if p.get("abstract") else "—",
            } for p in filtered[:200]])
            st.dataframe(df, hide_index=True, use_container_width=True, height=300)

            # Trend analysis
            if st.button("🧠 Analyze emerging trends with AI", type="primary", disabled=not client):
                if client:
                    with st.spinner(f"Analyzing {len(filtered)} papers for trends..."):
                        trends, err = analyze_journal_trends(filtered)
                        if trends:
                            st.session_state.trends_data = trends
                        else:
                            st.error(f"Error: {err}")

        if st.session_state.trends_data:
            st.divider()
            st.markdown("### 📊 Emerging trends")
            data = st.session_state.trends_data
            if data.get("summary"):
                st.info(data["summary"])
            for i, t in enumerate(data.get("trends", []), 1):
                with st.container(border=True):
                    st.markdown(f"**Trend {i} — {t.get('name','')}**")
                    st.markdown(t.get("definition", ""))
                    if t.get("supporting_papers"):
                        st.caption(f"Examples: {', '.join(t['supporting_papers'][:6])}")
                    if t.get("significance"):
                        st.caption(f"💡 {t['significance']}")

            # Optional: add these trends papers to corpus
            if st.button("➕ Add filtered papers to main corpus for screening"):
                added = 0
                existing_ids = {r.get("id") for r in st.session_state.results}
                for p in filtered:
                    if p.get("id") not in existing_ids:
                        st.session_state.results.append(p)
                        added += 1
                st.success(f"Added {added} new papers to search results. Switch to 'Screen' tab to triage them.")
